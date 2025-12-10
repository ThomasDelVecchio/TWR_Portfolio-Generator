import pandas as pd
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =====================================================================
# Formatting Helpers (FULL N/A FIX)
# =====================================================================

def fmt_pct_clean(x):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"{float(x)*100:.2f}%"
    except:

        return "N/A"

def fmt_dollar_clean(x):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"${float(x):,.2f}"
    except:
        return "N/A"

def safe(x):
    return "N/A" if x is None or pd.isna(x) else x

# =====================================================================
# Build a Light Grid Accent 1 table (your style)
# =====================================================================

def add_table(doc, headers, rows, right_align=None):

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ----- FORCE TABLE TO FIT PAGE WIDTH -----
    table.autofit = True
    table.allow_autofit = True

    # total table width target = ~6.2"
    max_width = Inches(6.2)
    col_width = max_width / len(headers)

    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # ----- HEADER -----
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    # ----- DATA ROWS -----
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

            if right_align and i in right_align:
                for p in cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Prevent row splitting across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)

    return table
