#!/usr/bin/env python3
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
from config import TARGET_PORTFOLIO_VALUE, TARGET_MONTHLY_CONTRIBUTION
from main1 import run_engine, fetch_price_history, load_cashflows_external, modified_dietz_for_ticker_window, get_portfolio_horizon_start
from io import BytesIO
import matplotlib
matplotlib.use("Agg")  # non-GUI backend; no windows

import matplotlib.pyplot as plt
import os
from config import ETF_SECTOR_MAP
from collections import defaultdict
from main1 import load_transactions_raw
from datetime import datetime

# ============================================================
# GLOBAL COLOR PALETTE FOR ALL CHARTS
# ============================================================
GLOBAL_PALETTE = [
    "#4C6A92",  # steel blue
    "#8C9CB1",  # soft gray-blue
    "#C0504D",  # muted red
    "#D79E9C",  # soft red-gray
    "#9BBB59",  # olive green
    "#C5D6A4",  # light olive
    "#8064A2",  # muted purple
    "#B1A0C7",  # lavender gray
    "#4F81BD",  # corporate blue
    "#A5B5CF",  # cool gray-blue
    "#F2C200",  # muted gold (accent)
    "#D6B656",  # soft gold-gray
]

import matplotlib as mpl

# Apply palette to ALL charts automatically
mpl.rcParams['axes.prop_cycle'] = mpl.cycler(color=GLOBAL_PALETTE)

# Optional consistent styling defaults
mpl.rcParams['figure.dpi'] = 150
mpl.rcParams['savefig.dpi'] = 150
mpl.rcParams['axes.titlesize'] = 12
mpl.rcParams['axes.labelsize'] = 11
mpl.rcParams['xtick.labelsize'] = 10
mpl.rcParams['ytick.labelsize'] = 10
mpl.rcParams['legend.fontsize'] = 10

# =====================================================================
# Formatting Helpers (FULL N/A FIX)
# =====================================================================

def fmt_pct_clean(x):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"{float(x)*100:.2f}%"
    except:

        return "N/A"

def fmt_dollar_clean(x):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"${float(x):,.2f}"
    except:
        return "N/A"

def safe(x):
    return "N/A" if x is None or pd.isna(x) else x

# =====================================================================
# Build a Light Grid Accent 1 table (your style)
# =====================================================================

def add_table(doc, headers, rows, right_align=None):

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ----- FORCE TABLE TO FIT PAGE WIDTH -----
    table.autofit = True
    table.allow_autofit = True

    # total table width target = ~6.2"
    max_width = Inches(6.2)
    col_width = max_width / len(headers)

    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # ----- HEADER -----
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    # ----- DATA ROWS -----
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

            if right_align and i in right_align:
                for p in cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Prevent row splitting across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)

    return table


# =====================================================================
# MAIN REPORT BUILDER — FULLY CORRECTED
# =====================================================================

def build_report():

    # Run the engine (unchanged math)
    twr_df, sec_full, class_full, pv, twr_si, twr_si_annualized, pl_si = run_engine()
    
    twr_raw = twr_df.copy()  # keep numeric returns for excess-return calc


    cf_ext = load_cashflows_external()

    tx_raw = load_transactions_raw()

    
    # === Inception date (same logic as run_engine) ===
    dates = []
    if not cf_ext.empty:
        dates.append(cf_ext["date"].min())
    if not tx_raw.empty:
        dates.append(tx_raw["date"].min())
    dates.append(pv.index.min())

    inception_date = min(dates)

    
    # External cashflows for proper P/L (deposits/withdrawals only)
    cf_ext = load_cashflows_external()

    # =============================================================
    # CLEAN ALL RETURN COLUMNS
    # =============================================================

    # Fix TWR DF
    twr_df["Return"] = twr_df["Return"].apply(fmt_pct_clean)

    # Fix Security-level DF
    if not sec_full.empty:
        for col in ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "YTD"]:
            if col in sec_full.columns:
                sec_full[col] = sec_full[col].apply(fmt_pct_clean)

    # =============================================================
    # START DOC
    # =============================================================

    doc = Document()

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    base.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # =============================================================
    # COVER PAGE
    # =============================================================
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
    doc.add_paragraph("\n\n\n\n")

    title = doc.add_paragraph("Portfolio Performance Report")
    title.runs[0].font.size = Pt(26)
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Time-Weighted & Money-Weighted Performance (Automated)")
    created_line = doc.add_paragraph(f"Created for Tom Short — {timestamp}")
    created_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    created_line.runs[0].font.size = Pt(12)
    created_line.runs[0].italic = True
    created_line.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(70, 130, 180)

    doc.add_page_break()

    # =============================================================
    # PORTFOLIO SNAPSHOT (VERTICAL FORMAT + REAL P/L)
    # =============================================================

    doc.add_heading("Portfolio Snapshot", level=1)
    doc.add_paragraph("Summary of time-weighted returns and P/L.")

    # Extract Return % values
    snap_map = {}
    for _, row in twr_df.iterrows():
        snap_map[row["Horizon"]] = row["Return"]

    # Horizons in EXACT order
    horizons = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y"]

    # REAL P/L calculator using pv and external cashflows
    # P/L = MV_end − MV_start − net_external_flows(start, end)
    
    def get_horizon_start(label: str) -> pd.Timestamp | None:
        """
        Thin wrapper around main1.get_portfolio_horizon_start so that
        horizon anchors for P/L, benchmark charts, etc. are ALWAYS
        identical to the ones used by compute_horizon_twr().
        """
        return get_portfolio_horizon_start(pv, inception_date, label)


    def compute_horizon_pl(h):
        """
        Portfolio P/L over horizon h using the SAME horizon start as TWR.
        P/L = MV_end − MV_start − net_external_flows(start, end)
        """
        as_of = pv.index.max()

        start = get_horizon_start(h)
        if start is None or start >= as_of:
            return "N/A"

        # ----- Map horizon start onto actual PV index -----
        if start not in pv.index:
            pv_idx = pv.index.sort_values()
            pos = pv_idx.searchsorted(start)
            if pos >= len(pv_idx):
                return "N/A"
            start = pv_idx[pos]

        mv_start = float(pv.loc[start])
        mv_end   = float(pv.loc[as_of])


        # flows strictly inside (start, as_of)
        net_flows = 0.0
        if cf_ext is not None and not cf_ext.empty:
            mask = (cf_ext["date"] > start) & (cf_ext["date"] < as_of)
            net_flows = float(cf_ext.loc[mask, "amount"].sum())

        pl = mv_end - mv_start - net_flows
        return fmt_dollar_clean(pl)

    # =============================================================
    # SINCE INCEPTION — PORTFOLIO RETURN & P/L
    # Uses SAME PV series, SAME inception_date, SAME external-flow logic.
    # =============================================================

    def compute_since_inception_pl():
        """Economic P/L since inception."""
        as_of = pv.index.max()
        start = inception_date

        # Map inception_date onto nearest PV trading day
        if start not in pv.index:
            pv_idx = pv.index.sort_values()
            pos = pv_idx.searchsorted(start)
            if pos >= len(pv_idx):
                return "N/A"
            start = pv_idx[pos]

        mv_start = float(pv.loc[start])
        mv_end   = float(pv.loc[as_of])


        # external flows strictly AFTER inception (not including day 1 capital)
        net_flows = 0.0
        if cf_ext is not None and not cf_ext.empty:
            mask = (cf_ext["date"] > start) & (cf_ext["date"] < as_of)
            net_flows = float(cf_ext.loc[mask, "amount"].sum())

        pl = mv_end - mv_start - net_flows
        return pl


    def compute_since_inception_return():
        """
        Wrapper to expose the since-inception portfolio TWR that was
        already computed in run_engine(). This avoids maintaining a
        second TWR aggregation loop in this file.
        """
        return twr_si


    # Build vertical snapshot rows
    rows = []
    for h in horizons:
        ret = snap_map.get(h, "N/A")
        pl = compute_horizon_pl(h) if ret != "N/A" else "N/A"
        rows.append([h, ret, pl])
        
    # ===== ADD SINCE-INCEPTION ROW =====
    si_pl = pl_si

    # Prefer annualized SI if run_engine computed it (i.e., inception > ~1 year).
    if twr_si_annualized is not None and not pd.isna(twr_si_annualized):
        si_ret_for_display = twr_si_annualized
    else:
        si_ret_for_display = twr_si

    rows.append([
        "Since Inception",
        fmt_pct_clean(si_ret_for_display),
        fmt_dollar_clean(si_pl)
    ])



    add_table(
        doc,
        ["Horizon", "Return %", "P/L ($)"],
        rows,
        right_align=[1, 2]
    )

    # Small spacing before summary table
    doc.add_paragraph().paragraph_format.space_before = Pt(2)

    # Portfolio Summary Table
    total_value = pv.iloc[-1] if isinstance(pv, pd.Series) else pv
    num_holdings = len(sec_full)

    summary_rows = [
        ["Total Value", fmt_dollar_clean(pv.iloc[-1])],
        ["Target Portfolio Value", fmt_dollar_clean(TARGET_PORTFOLIO_VALUE)],
        ["Number of Holdings", num_holdings],
    ]

    add_table(
        doc,
        ["Metric", "Value"],
        summary_rows,
        right_align=[1],
    )

    # PREP FOR PERFORMANCE HIGHLIGHTS — BUILD sec_only + prices

    holdings_df = pd.read_csv("sample holdings.csv")
    holdings_df.columns = [c.lower() for c in holdings_df.columns]
    holdings_df["ticker"] = holdings_df["ticker"].str.upper()
    if "target_pct" not in holdings_df.columns:
        holdings_df["target_pct"] = 0.0
    else:
        holdings_df["target_pct"] = holdings_df["target_pct"].astype(float)

    # Merge target_pct into sec_full (include CASH now)
    sec_only = sec_full.copy()
    sec_only = sec_only.merge(
        holdings_df[["ticker", "target_pct"]],
        on="ticker",
        how="left"
    )
    sec_only["target_pct"] = sec_only["target_pct"].fillna(0.0)

    # Shorten asset class names
    asset_class_map = {
        "US Large Cap": "US LC",
        "US Growth": "US Growth",
        "US Small Cap": "US SC",
        "International Equity": "INTL EQTY",
        "Gold / Precious Metals": "GOLD / PM",
        "Digital Assets": "DIGITAL",
        "US Bonds": "US Bonds",
        "CASH": "Cash"
    }
    sec_only["asset_class_short"] = sec_only["asset_class"].map(lambda x: asset_class_map.get(x, x))

    # Fetch live prices for non-cash tickers
    tickers = sec_only[sec_only["ticker"] != "CASH"]["ticker"].tolist()
    prices = fetch_price_history(tickers)
    latest_prices = prices.iloc[-1]

    # Helper: compute P/L for a ticker inside a horizon window
    def compute_ticker_pl(ticker, h):
        """
        Correct economic P/L for a single ticker over a horizon.

        P/L = MV_end – MV_start – net_internal_flows(start, end)

        SPECIAL CASE:
          - For h == "SI": horizon is from this ticker's FIRST TRADE DATE
            (not portfolio inception). This guarantees a non-NaN SI for any
            ticker that has at least one transaction and price history.
        """
        if ticker == "CASH":
            # Treat CASH as 0% return, 0 P/L for horizons in this table.
            return fmt_dollar_clean(0.0)

        # price series
        if ticker not in prices.columns:
            return "N/A"
        series = prices[ticker].dropna()
        if series.empty:
            return "N/A"

        as_of_port = pv.index.max()
        as_of_price = series.index.max()
        as_of = min(as_of_port, as_of_price)

        # ----- Load transactions for this ticker -----
        tx = load_transactions_raw().copy()
        tx = tx[tx["ticker"] == ticker].copy()
        tx = tx.sort_values("date")

        if tx.empty:
            return "N/A"

        first_trade = tx["date"].min()

        # =================================================================
        # SI: since *this ticker's* inception (first trade), not portfolio
        # =================================================================
        if h == "SI":
            # Start at later of first trade and earliest price in the series
            earliest_px = series.index.min()
            raw_start = max(first_trade, earliest_px)

            # Snap to the first available price ON or AFTER raw_start
            series_dates = series.index.sort_values()
            idx = series_dates.searchsorted(raw_start)
            if idx >= len(series_dates):
                return "N/A"

            start = series_dates[idx]

            if start >= as_of:
                return "N/A"


        else:
            # =============================================================
            # ORIGINAL HORIZON LOGIC (UNTOUCHED FOR NON-SI HORIZONS)
            # =============================================================
            raw_start = get_horizon_start(h)
            if raw_start is None or raw_start >= as_of:
                return "N/A"

            # Clamp to earliest price date for this ticker
            earliest_px = series.index.min()
            if raw_start < earliest_px:
                raw_start = earliest_px

            series_dates = series.index.sort_values()

            # MTD uses same "last prior price" logic
            if h == "MTD":
                # raw_start = prior-month-end; use FIRST price *after* that date
                idx = series_dates.searchsorted(raw_start)
                if idx >= len(series_dates):
                    return "N/A"
                start = series_dates[idx]

            # 1D MUST use strict previous trading day only
            elif h == "1D":
                # 1D should match the portfolio horizon exactly: raw_start is the
                # previous trading day at the portfolio level, so we want the
                # first price on or AFTER raw_start for this ticker.
                idx = series_dates.searchsorted(raw_start)
                if idx >= len(series_dates):
                    return "N/A"
                start = series_dates[idx]

            # All other horizons: nearest prior price
            else:
                idx = series_dates.searchsorted(raw_start, side="right") - 1
                if idx < 0:
                    return "N/A"
                start = series_dates[idx]

            if start >= as_of:
                return "N/A"

            # Not owned at start → no P/L
            if first_trade > start:
                return "N/A"

            # Horizon must not start before first trade
            start = max(start, first_trade)

        # ----- Prices -----
        try:
            px_start = float(series.loc[start])
            px_end = float(series.loc[as_of])
        except Exception:
            return "N/A"

        # ----- Shares at end -----
        row = sec_only[sec_only["ticker"] == ticker]
        if row.empty:
            return "N/A"
        shares_end = float(row["shares"].iloc[0])

        # ----- Shares at start -----
        mask = tx["date"] <= start
        shares_start = tx.loc[mask, "shares"].sum() if mask.any() else 0.0

        # ----- Internal flows inside window -----
        mask2 = (tx["date"] > start) & (tx["date"] < as_of)
        # Our file uses amount negative for buys (cash out), positive for sells (cash in).
        # When computing economic P/L, we subtract net internal flows (same as before).
        net_internal = -tx.loc[mask2, "amount"].sum()

        # ----- Economic P/L -----
        mv_start = shares_start * px_start
        mv_end = shares_end * px_end

        pl = mv_end - mv_start - net_internal

        return fmt_dollar_clean(pl)


    # ---------------------------------------------------------------
    # PERFORMANCE HIGHLIGHTS TABLE
    # ---------------------------------------------------------------
    doc.add_heading("Performance Highlights", level=1)

    perf_rows = []

    if not sec_full.empty:
        # Build numeric version for selecting top/bottom performers
        sec_full_numeric = sec_full.copy()
        for col in ["1M", "1D"]:
            if col in sec_full_numeric.columns:
                sec_full_numeric[col] = (
                    sec_full_numeric[col].astype(str)
                    .str.replace("%", "", regex=False)
                    .replace("N/A", np.nan)
                    .astype(float)
                )

        # ---------------- TOP / BOTTOM 1M ----------------
        if "1M" in sec_full_numeric and not sec_full_numeric["1M"].dropna().empty:
            top_1m = sec_full_numeric.loc[sec_full_numeric["1M"].idxmax()]
            bottom_1m = sec_full_numeric.loc[sec_full_numeric["1M"].idxmin()]

            # Correct Ticker-Level P/L using your function
            top_pl_1m = compute_ticker_pl(top_1m["ticker"], "1M")
            bottom_pl_1m = compute_ticker_pl(bottom_1m["ticker"], "1M")

            perf_rows.append([
                "Top 1M Performer",
                f"{top_1m['ticker']} ({fmt_pct_clean(top_1m['1M']/100)}, {top_pl_1m})"
            ])
            perf_rows.append([
                "Bottom 1M Performer",
                f"{bottom_1m['ticker']} ({fmt_pct_clean(bottom_1m['1M']/100)}, {bottom_pl_1m})"
            ])
        else:
            perf_rows.append(["Top 1M Performer", "N/A"])
            perf_rows.append(["Bottom 1M Performer", "N/A"])

        # ---------------- BEST / WORST 1D ----------------
        if "1D" in sec_full_numeric and not sec_full_numeric["1D"].dropna().empty:
            best_1d = sec_full_numeric.loc[sec_full_numeric["1D"].idxmax()]
            bottom_1d = sec_full_numeric.loc[sec_full_numeric["1D"].idxmin()]

            # Correct Ticker-Level P/L using your function
            best_pl_1d = compute_ticker_pl(best_1d["ticker"], "1D")
            bottom_pl_1d = compute_ticker_pl(bottom_1d["ticker"], "1D")

            perf_rows.append([
                "Best 1D Performer",
                f"{best_1d['ticker']} ({fmt_pct_clean(best_1d['1D']/100)}, {best_pl_1d})"
            ])
            perf_rows.append([
                "Bottom 1D Performer",
                f"{bottom_1d['ticker']} ({fmt_pct_clean(bottom_1d['1D']/100)}, {bottom_pl_1d})"
            ])
        else:
            perf_rows.append(["Best 1D Performer", "N/A"])
            perf_rows.append(["Bottom 1D Performer", "N/A"])

    else:
        perf_rows = [
            ["Top 1M Performer", "N/A"],
            ["Bottom 1M Performer", "N/A"],
            ["Best 1D Performer", "N/A"],
            ["Bottom 1D Performer", "N/A"],
        ]

    # Add table
    add_table(
        doc,
        ["Metric", "Value"],
        perf_rows,
        right_align=[1]
    )


    # ---------------------------------------------------------------
    # RISK & DIVERSIFICATION TABLE
    # ---------------------------------------------------------------
    doc.add_heading("Risk & Diversification", level=1)

    # Exclude CASH from weight calculations
    sec_no_cash = sec_full[sec_full["ticker"] != "CASH"].copy()

    # Top 3 holdings % of portfolio
    top3_pct = sec_no_cash.nlargest(3, "weight")["weight"].sum() * 100 if not sec_no_cash.empty else 0

    # Compute total weights per asset class
    asset_class_weights = sec_no_cash.groupby("asset_class")["weight"].sum() * 100

    # Largest asset class by current weight
    largest_class = asset_class_weights.idxmax() if not asset_class_weights.empty else "N/A"
    largest_class_pct = asset_class_weights.max() if not asset_class_weights.empty else 0

    # Compute target percentages from sample holdings
    holdings_df = pd.read_csv("sample holdings.csv")
    target_pct_map = holdings_df.groupby("asset_class")["target_pct"].sum().to_dict()

    # Largest over/underweight relative to target
    largest_over = None
    largest_under = None
    max_diff = -np.inf
    min_diff = np.inf

    for ac, wt in asset_class_weights.items():
        target = target_pct_map.get(ac, 0)
        diff = wt - target
        if diff > max_diff:
            max_diff = diff
            largest_over = f"{ac} ({wt:.2f}% vs {target:.2f}%)"
        if diff < min_diff:
            min_diff = diff
            largest_under = f"{ac} ({wt:.2f}% vs {target:.2f}%)"

    risk_rows = [
        ["Top 3 holdings % of portfolio", f"{top3_pct:.2f}%"],
        ["Largest asset class", f"{largest_class} ({largest_class_pct:.2f}%)"],
        ["Largest overweight", largest_over if largest_over else "N/A"],
        ["Largest underweight", largest_under if largest_under else "N/A"],
    ]

    add_table(
        doc,
        ["Metric", "Value"],
        risk_rows,
        right_align=[1]
    )
   

    # ---------------------------------------------------------------
    # FLOWS SUMMARY (Unified External + Internal) — YTD VERSION
    # ---------------------------------------------------------------
    doc.add_heading("Flows Summary (YTD)", level=1)

    as_of = pv.index.max()
    ytd_start = as_of.replace(month=1, day=1)

    # External flows (deposits/withdrawals)
    flows_ext = cf_ext.copy()
    flows_ext = flows_ext[flows_ext["date"] >= ytd_start].sort_values("date")

    if flows_ext.empty:
        ytd_deposits = 0.0
        ytd_withdrawals = 0.0
        net_ytd_ext = 0.0
        most_recent_ext = None
    else:
        ytd_deposits = flows_ext.loc[flows_ext["amount"] > 0, "amount"].sum()
        ytd_withdrawals = flows_ext.loc[flows_ext["amount"] < 0, "amount"].sum()
        net_ytd_ext = flows_ext["amount"].sum()
        most_recent_ext = flows_ext["date"].max()

    # Internal trades (buys/sells)
    tx_raw = load_transactions_raw().copy()
    tx_raw = tx_raw[tx_raw["date"] >= ytd_start].sort_values("date")

    if tx_raw.empty:
        ytd_buys = 0.0
        ytd_sells = 0.0
        net_ytd_internal = 0.0
        most_recent_tx = None
    else:
        # buys = negative amounts (cash out)
        ytd_buys = tx_raw.loc[tx_raw["amount"] < 0, "amount"].sum()
        ytd_sells = tx_raw.loc[tx_raw["amount"] > 0, "amount"].sum()
        net_ytd_internal = ytd_buys + ytd_sells
        most_recent_tx = tx_raw["date"].max()

    # Choose the most recent of ANY flow
    if most_recent_ext and most_recent_tx:
        most_recent_any = max(most_recent_ext, most_recent_tx).strftime("%Y-%m-%d")
    elif most_recent_ext:
        most_recent_any = most_recent_ext.strftime("%Y-%m-%d")
    elif most_recent_tx:
        most_recent_any = most_recent_tx.strftime("%Y-%m-%d")
    else:
        most_recent_any = "N/A"

    flow_rows = [
        ["YTD Net External Flows", fmt_dollar_clean(net_ytd_ext)],
        ["• YTD Deposits", fmt_dollar_clean(ytd_deposits)],
        ["• YTD Withdrawals", fmt_dollar_clean(ytd_withdrawals)],
        ["YTD Net Internal Trading Flows", fmt_dollar_clean(net_ytd_internal)],
        ["• YTD Buys (Cash Out)", fmt_dollar_clean(ytd_buys)],
        ["• YTD Sells (Cash In)", fmt_dollar_clean(ytd_sells)],
        ["Most Recent Flow", most_recent_any],
    ]

    add_table(
        doc,
        ["Metric", "Value"],
        flow_rows,
        right_align=[1]
    )



    # ---------------------------------------------------------------
    # PORTFOLIO COMPOSITION & STRATEGY TABLE
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Portfolio Composition & Strategy", level=1)
    doc.add_heading("Holdings by Ticker", level=2)

    # Compute prices (CASH = 1)
    sec_only["price"] = sec_only["ticker"].map(lambda t: 1 if t == "CASH" else latest_prices[t])
    sec_only["value"] = sec_only["shares"] * sec_only["price"]
    total_value = sec_only["value"].sum()
    sec_only["allocation"] = sec_only["value"] / total_value * 100

    # Compute "To Contrib" = amount needed to reach target weight, never negative
    sec_only["to_contrib"] = np.maximum(((sec_only["target_pct"] - sec_only["allocation"]) / 100) * total_value, 0)

    # Format columns for display
    sec_only["allocation"] = sec_only["allocation"].map(lambda x: f"+{x:.2f}%" if x >= 0 else f"{x:.2f}%")
    sec_only["target_pct"] = sec_only["target_pct"].map(lambda x: f"+{x:.2f}%" if x >= 0 else f"{x:.2f}%")
    sec_only["to_contrib"] = sec_only["to_contrib"].map(lambda x: f"${x:,.2f}")

    # Build table rows
    table_rows = []
    for _, row in sec_only.iterrows():
        table_rows.append([
            row["ticker"],
            row.get("asset_class_short", "Unknown"),
            f"{row['shares']:.3f}",
            fmt_dollar_clean(row["price"]),
            fmt_dollar_clean(row["value"]),
            row["allocation"],
            row["target_pct"],
            row["to_contrib"],
        ])

    # Add TOTAL row
    total_value_sum = sec_only["value"].sum()
    total_to_contrib_sum = sec_only["to_contrib"].replace(r"[\$,]", "", regex=True).astype(float).sum()
    table_rows.append([
        "TOTAL",
        "",
        "",
        "",
        fmt_dollar_clean(total_value_sum),
        "",
        "",
        f"${total_to_contrib_sum:,.2f}"
    ])

    # Add table
    add_table(
        doc,
        ["Ticker", "Asset Class", "Shares", "Price ($)", "Value ($)", "Allocation", "Target", "To Contrib"],
        table_rows,
        right_align=[2,3,4,5,6,7]
    )

    # ---------------------------------------------------------------
    # ILLUSTRATIVE MONTHLY CONTRIBUTION SCHEDULE
    # ---------------------------------------------------------------
    doc.add_heading("Illustrative Monthly Contribution Schedule", level=1)

    # Filter out holdings with zero to_contrib
    monthly_df = sec_only.copy()
    monthly_df["to_contrib_numeric"] = monthly_df["to_contrib"].replace(r"[\$,]", "", regex=True).astype(float)
    monthly_df = monthly_df[monthly_df["to_contrib_numeric"] > 0].copy()

    if not monthly_df.empty:
        # Use configurable monthly contribution from config
        total_monthly = TARGET_MONTHLY_CONTRIBUTION  # import this from config
        total_gap = monthly_df["to_contrib_numeric"].sum()
        monthly_df["monthly_contrib"] = monthly_df["to_contrib_numeric"] / total_gap * total_monthly
        monthly_df["share_of_monthly"] = monthly_df["monthly_contrib"] / total_monthly * 100

        # Build table rows
        rows = []
        for _, r in monthly_df.iterrows():
            rows.append([
                r["ticker"],
                r.get("asset_class_short", "Unknown"),
                fmt_dollar_clean(r["to_contrib_numeric"]),
                fmt_dollar_clean(r["monthly_contrib"]),
                f"{r['share_of_monthly']:.1f}%",
            ])

        add_table(
            doc,
            ["Ticker", "Asset Class", "Gap to Target", "Monthly Contrib", "Share of Monthly"],
            rows,
            right_align=[2,3,4]
        )

        # Footer paragraph
        footer = ("At approximately "
                  f"${total_monthly:,.0f}/month, this schedule allocates contributions "
                  "proportionally to each holding's gap. It would take about "
                  f"{total_gap / total_monthly:.1f} months to close all gaps, assuming flat markets.")
        doc.add_paragraph(footer)

    # ---------------------------------------------------------------
    # ASSET CLASS ALLOCATION TABLE
    # ---------------------------------------------------------------
    doc.add_heading("Asset Class Allocation", level=1)

    # Load sample holdings for target_pct (including CASH if needed)
    holdings_df = pd.read_csv("sample holdings.csv")
    holdings_df.columns = [c.lower() for c in holdings_df.columns]
    holdings_df["ticker"] = holdings_df["ticker"].str.upper()
    holdings_df["target_pct"] = holdings_df.get("target_pct", 0.0).astype(float)

    # Exclude CASH for asset class calculations if you want
    sec_no_cash = sec_only.copy()  # sec_only already includes CASH if needed
    total_value = sec_no_cash["value"].sum()

    # Merge short asset class and target_pct from holdings
    sec_merge = sec_no_cash.merge(
        holdings_df[["ticker", "target_pct"]],
        on="ticker",
        how="left",
        suffixes=("", "_holdings")
    )

    # Compute actual allocations per asset class
    asset_group = (
        sec_merge.groupby("asset_class_short")
        .agg(
            value=("value", "sum"),
            target_pct=("target_pct_holdings", "sum")
        )
        .reset_index()
    )

    # Compute actual percentage allocation
    asset_group["actual_pct"] = asset_group["value"] / total_value * 100

    # Compute delta
    asset_group["delta_pct"] = asset_group["actual_pct"] - asset_group["target_pct"]

    # Format columns for display
    asset_group["actual_pct_fmt"] = asset_group["actual_pct"].map(lambda x: f"{x:.2f}%")
    asset_group["target_pct_fmt"] = asset_group["target_pct"].map(lambda x: f"{x:.2f}%")
    asset_group["delta_pct_fmt"] = asset_group["delta_pct"].map(lambda x: f"+{x:.2f}%" if x >= 0 else f"{x:.2f}%")

    # Build table rows
    table_rows = []
    for _, row in asset_group.iterrows():
        table_rows.append([
            row["asset_class_short"],
            fmt_dollar_clean(row["value"]),
            row["actual_pct_fmt"],
            row["target_pct_fmt"],
            row["delta_pct_fmt"]
        ])

    # Add TOTAL row
    total_value_sum = asset_group["value"].sum()
    total_actual_pct = asset_group["actual_pct"].sum()
    total_target_pct = asset_group["target_pct"].sum()
    total_delta_pct = asset_group["delta_pct"].sum()

    table_rows.append([
        "TOTAL",
        fmt_dollar_clean(total_value_sum),
        f"{total_actual_pct:.2f}%",
        f"{total_target_pct:.2f}%",
        f"{total_delta_pct:+.2f}%"
    ])

    # Add table
    add_table(
        doc,
        ["Asset Class", "Value ($)", "Actual %", "Target %", "Delta %"],
        table_rows,
        right_align=[1, 2, 3, 4]
    )

    # Footer
    doc.add_paragraph("Delta % = actual allocation minus target allocation").paragraph_format.space_before = Pt(2)

    doc.add_page_break()


    # =============================================================
    # TICKER ALLOCATION PIE CHART
    # =============================================================
    # Build ticker allocation group
    ticker_group = sec_only.copy()
    ticker_group = ticker_group[ticker_group["value"] > 0]  # exclude zero-value tickers

    # Sort descending by value to avoid adjacent small slices
    ticker_group = ticker_group.sort_values(by="value", ascending=False)

    ticker_labels = ticker_group["ticker"].tolist()
    ticker_values = ticker_group["value"].tolist()

    fig, ax = plt.subplots(figsize=(6, 6))
    wedges, texts, autotexts = ax.pie(
        ticker_values,
        labels=ticker_labels,
        autopct="%1.1f%%",
        startangle=90,
        wedgeprops={"edgecolor": "w"},
        pctdistance=0.75,  # move % labels outward to reduce overlap
        labeldistance=1.05  # labels slightly outside
    )
    ax.axis("equal")  # make circle

    # Style text
    plt.setp(texts, size=8)
    plt.setp(autotexts, size=7, weight="bold", color="black")

    # Save chart to in-memory PNG
    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    img_stream.seek(0)

    # Insert into DOCX
    doc.add_heading("Ticker Allocation", level=1)

    # Create a paragraph for the image and center it
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(5))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Caption centered
    caption = doc.add_paragraph("Pie chart of current ticker allocations.", style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ---------------------------------------------------------------
    # TICKER ALLOCATION VS TARGET BAR CHART
    # ---------------------------------------------------------------

    # Get path to sample holdings in same folder as this script
    script_dir = os.path.dirname(os.path.abspath(__file__))
    holdings_path = os.path.join(script_dir, "sample holdings.csv")

    holdings_df = pd.read_csv(holdings_path)
    holdings_df.columns = [c.lower() for c in holdings_df.columns]
    holdings_df["ticker"] = holdings_df["ticker"].str.upper()
    if "target_pct" not in holdings_df.columns:
        holdings_df["target_pct"] = 0.0
    else:
        holdings_df["target_pct"] = holdings_df["target_pct"].astype(float)

    # Merge actual values from sec_only with target_pct
    ticker_merge = sec_only[["ticker", "value"]].merge(
        holdings_df[["ticker", "target_pct"]],
        on="ticker",
        how="left"
    )
    ticker_merge["target_pct"] = ticker_merge["target_pct"].fillna(0.0)

    # Compute actual % allocation
    total_value = ticker_merge["value"].sum()
    ticker_merge["actual_pct"] = ticker_merge["value"] / total_value * 100

    # Plot side-by-side bar chart
    fig, ax = plt.subplots(figsize=(10, 5))
    width = 0.35
    x = np.arange(len(ticker_merge))
    bars1 = ax.bar(x - width/2, ticker_merge["actual_pct"], width, label="Actual %")
    bars2 = ax.bar(x + width/2, ticker_merge["target_pct"], width, label="Target %")

    # Add text labels above bars
    for i in range(len(ticker_merge)):
        ax.text(i - width/2, ticker_merge["actual_pct"].iloc[i] + 0.5,
                f"{ticker_merge['actual_pct'].iloc[i]:.1f}%", ha='center', va='bottom', fontsize=8)
        ax.text(i + width/2, ticker_merge["target_pct"].iloc[i] + 0.5,
                f"{ticker_merge['target_pct'].iloc[i]:.1f}%", ha='center', va='bottom', fontsize=8)

    # X-axis
    ax.set_xticks(x)
    ax.set_xticklabels(ticker_merge["ticker"], rotation=45, ha="right", fontsize=9)
    ax.set_ylabel("Allocation (%)")
    ax.set_title("")
    ax.legend()
    plt.tight_layout()

    # Save chart to in-memory PNG
    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    img_stream.seek(0)

    # Insert into DOCX and center
    doc.add_heading("Ticker Allocation vs Target", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Actual vs target allocation by ticker.", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ---------------------------------------------------------------
    # ASSET CLASS ALLOCATION PIE & VS TARGET
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Asset Class Allocation Breakdown", level=1)

    # Load sample holdings dynamically from the same folder as the script
    script_dir = os.path.dirname(os.path.abspath(__file__))
    holdings_path = os.path.join(script_dir, "sample holdings.csv")

    holdings_df = pd.read_csv(holdings_path)
    holdings_df.columns = [c.lower() for c in holdings_df.columns]
    holdings_df["ticker"] = holdings_df["ticker"].str.upper()
    holdings_df["target_pct"] = holdings_df.get("target_pct", 0.0).astype(float)

    # Merge actual values from sec_only
    sec_only_grouped = (
        sec_only.groupby("asset_class_short")
        .agg(value=("value", "sum"))
        .reset_index()
    )
    merged = holdings_df.groupby("asset_class")["target_pct"].sum().reset_index()
    merged["asset_class_short"] = merged["asset_class"].map({
        "US Large Cap":"US LC",
        "US Growth":"US Growth",
        "US Small Cap":"US SC",
        "International Equity":"INTL EQTY",
        "Gold / Precious Metals":"GOLD / PM",
        "Digital Assets":"DIGITAL",
        "US Bonds":"US Bonds",
        "CASH":"CASH"
    })
    ticker_merge = pd.merge(sec_only_grouped, merged[["asset_class_short","target_pct"]],
                            on="asset_class_short", how="left").fillna(0.0)

    # Sort descending by actual value
    ticker_merge = ticker_merge.sort_values("value", ascending=False).reset_index(drop=True)
    ticker_merge["actual_pct"] = ticker_merge["value"] / ticker_merge["value"].sum() * 100

    # ---------------- PIE CHART ----------------
    fig, ax = plt.subplots(figsize=(6,6))  # same size as ticker pie
    wedges, texts, autotexts = ax.pie(
        ticker_merge["actual_pct"],
        labels=ticker_merge["asset_class_short"],
        autopct="%1.1f%%",
        startangle=90,
        wedgeprops={"edgecolor":"w"},
        pctdistance=0.75,   # same as ticker pie
        labeldistance=1.05  # same as ticker pie
    )
    plt.setp(texts, size=8)          # match ticker pie
    plt.setp(autotexts, size=7, weight="bold", color="black")  # match ticker pie
    ax.axis("equal")  # circle

    # Save chart to in-memory PNG
    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    img_stream.seek(0)

    # Insert into DOCX and center
    doc.add_heading("Asset Class Allocation", level=2)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(5))  # match ticker pie width
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ---------------- BAR CHART ----------------
    fig, ax = plt.subplots(figsize=(8,5))
    width = 0.35
    x = range(len(ticker_merge))
    ax.bar(x, ticker_merge["actual_pct"], width, label="Actual %")
    ax.bar([i + width for i in x], ticker_merge["target_pct"], width, label="Target %")

    # Percentages on top
    for i in x:
        ax.text(i, ticker_merge["actual_pct"].iloc[i]+0.5,
                f"{ticker_merge['actual_pct'].iloc[i]:.1f}%", ha="center", va="bottom", fontsize=9)
        ax.text(i + width, ticker_merge["target_pct"].iloc[i]+0.5,
                f"{ticker_merge['target_pct'].iloc[i]:.1f}%", ha="center", va="bottom", fontsize=9)

    ax.set_xticks([i + width/2 for i in x])
    ax.set_xticklabels(ticker_merge["asset_class_short"], rotation=45, ha="right", fontsize=10)
    ax.set_ylabel("Allocation (%)")
    ax.set_title("")
    ax.legend()

    img_stream = BytesIO()
    plt.tight_layout()
    plt.savefig(img_stream, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    img_stream.seek(0)

    doc.add_heading("Asset Class Allocation vs Target", level=2)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # ASSET ALLOCATION OVER TIME — STACKED AREA CHART
    # ---------------------------------------------------------------
    doc.add_heading("Asset Allocation Over Time", level=2)

    try:
        tx_hist = load_transactions_raw().copy()

        # Need both transactions and prices to build history
        if tx_hist.empty or prices.empty:
            doc.add_paragraph(
                "Insufficient transaction or price history to build an allocation-over-time chart."
            )
        else:

            # ---------------------------------------------------------------
            # Build a continuous daily index & forward-fill PV
            # ---------------------------------------------------------------
            full_index = pd.date_range(
                start=pv.index.min(),
                end=pv.index.max(),
                freq="D"
            )

            pv_mod = pv.reindex(full_index).ffill()
            alloc_index = full_index

            # Pivot transactions into daily share changes per ticker
            tx_hist["date"] = pd.to_datetime(tx_hist["date"])
            pos_changes = (
                tx_hist
                .pivot_table(
                    index="date",
                    columns="ticker",
                    values="shares",
                    aggfunc="sum"
                )
                .sort_index()
            )

            # Align positions to full daily index, FF to avoid dips
            pos_changes = pos_changes.reindex(alloc_index, fill_value=0.0)
            pos_daily = pos_changes.cumsum().ffill().bfill()

            # -------------------------------------------------------
            # Reconcile position history to current holdings
            # -------------------------------------------------------
            shares_map = (
                sec_only[["ticker", "shares"]]
                .set_index("ticker")["shares"]
                .to_dict()
            )

            # Adjust existing ticker columns to match final snapshot
            for t in list(pos_daily.columns):
                if t == "CASH":
                    continue
                if t in shares_map:
                    target = float(shares_map[t])
                    current = float(pos_daily[t].iloc[-1])
                    diff = target - current
                    if abs(diff) > 1e-8:
                        pos_daily[t] = pos_daily[t] + diff

            # Add static columns for tickers never seen in tx_hist
            for t, shares in shares_map.items():
                if t == "CASH":
                    continue
                if t not in pos_daily.columns and t in prices.columns:
                    pos_daily[t] = float(shares)

            # Restrict to tickers we actually have prices for
            common_tickers = [t for t in pos_daily.columns if t in prices.columns]

            if not common_tickers:
                doc.add_paragraph(
                    "No overlapping tickers between transactions and price history."
                )
            else:
                pos_daily = pos_daily[common_tickers]

                # Forward-fill prices on full daily index
                px_aligned = (
                    prices[common_tickers]
                    .reindex(alloc_index)
                    .ffill()
                    .bfill()
                )

                # Daily market value per ticker
                mv_daily = pos_daily * px_aligned

                # Map tickers → asset classes
                holdings_map = (
                    holdings_df[["ticker", "asset_class"]]
                    .drop_duplicates()
                    .set_index("ticker")["asset_class"]
                    .to_dict()
                )

                def map_asset_class_short(ticker: str) -> str:
                    full = holdings_map.get(ticker, "Unknown")
                    return asset_class_map.get(full, full)

                mv_daily.columns = [map_asset_class_short(t) for t in mv_daily.columns]

                # Aggregate by asset class
                mv_by_class = mv_daily.groupby(axis=1, level=0).sum()

                # Align PV
                pv_mod_aligned = pv_mod.reindex(mv_by_class.index).ffill().bfill()
                invested_total = mv_by_class.sum(axis=1)
                cash_series = pv_mod_aligned - invested_total
                mv_by_class["Cash"] = cash_series

                # Convert to allocation percentages
                total_mv = mv_by_class.sum(axis=1).replace(0, np.nan)
                alloc = mv_by_class.div(total_mv, axis=0).dropna(how="all")

                if alloc.empty:
                    doc.add_paragraph(
                        "No valid allocation history could be computed."
                    )
                else:
                    alloc_pct = alloc * 100.0

                    # ----------------------------------------------------------
                    # VISUAL-ONLY SMOOTHING (preserve row sum = 100%)
                    # ----------------------------------------------------------
                    kernel = np.array([0.25, 0.5, 0.25])

                    smooth_mat = []
                    for col in alloc_pct.columns:
                        series = alloc_pct[col].values
                        s = np.convolve(series, kernel, mode="same")
                        smooth_mat.append(s)

                    # Stack: (num_classes × num_days)
                    smooth_arr = np.vstack(smooth_mat)

                    # Renormalize to 100% exactly
                    row_sums = smooth_arr.sum(axis=0)
                    row_sums[row_sums == 0] = np.nan
                    smooth_arr = (smooth_arr / row_sums) * 100.0

                    # ----------------------------------------------------------
                    # PLOT — Bigger, cleaner, softer
                    # ----------------------------------------------------------
                    fig, ax = plt.subplots(figsize=(14, 9))

                    ax.stackplot(
                        alloc_pct.index,
                        smooth_arr,
                        labels=list(alloc_pct.columns),
                        alpha=0.92,
                        linewidth=0,
                        antialiased=True
                    )


                    FONT_SCALE = 1.5

                    ax.set_title("")
                    fig.suptitle("")
                    ax.set_ylabel(
                        "Allocation (%)",
                        fontsize=int(11 * FONT_SCALE)   
                    )

                    ax.tick_params(
                        axis="both",
                        labelsize=int(10 * FONT_SCALE)   
                    )

                    legend = ax.legend(
                        loc="upper left",
                        fontsize=int(10 * FONT_SCALE),    
                        ncol=3
                    )
                    for txt in legend.get_texts():
                        txt.set_fontsize(int(10 * FONT_SCALE))

                    fig.autofmt_xdate()

                    fig.tight_layout()


                    img_stream = BytesIO()
                    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
                    plt.close(fig)
                    img_stream.seek(0)

                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(7))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cap = doc.add_paragraph(
                        "Figure: Daily asset class allocation over time, including cash.",
                        style="Normal",
                    )
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception:
        doc.add_paragraph(
            "Asset allocation over time chart could not be generated.",
            style="Normal"
        )


    # ---------------------------------------------------------------
    # PORTFOLIO VS BENCHMARKS — MTD & YTD (PRICE RETURN)
    # ---------------------------------------------------------------

    # Keep ORIGINAL pv for horizon logic (TWR + MD + snapshot consistency)
    pv_trading = pv.sort_index()

    # Build a DAILY PV series for CHARTS ONLY
    pv_daily = pv_trading.reindex(
        pd.date_range(pv_trading.index.min(), pv_trading.index.max(), freq="D")
    ).ffill()

    portfolio_value = pv_daily.copy()



    # Current "as of" date
    as_of = portfolio_value.index[-1]

    # ---------------- MTD (Institutional: last trading day of prior month) ----------------
    pv_dates = pv.index
    mtd_start = get_horizon_start("MTD")
    if mtd_start is None:
        mtd_start = pv.index.min()

    # unified anchor for EVERYTHING (portfolio + benchmarks)
    pv_nonzero = pv[pv > 0]
    twr_anchor = max(mtd_start, pv_nonzero.index.min())

    portfolio_mtd = portfolio_value[portfolio_value.index >= twr_anchor]

    portfolio_mtd = (portfolio_mtd / portfolio_mtd.iloc[0] - 1) * 100


    # ---------------- YTD ----------------
    ytd_start = get_horizon_start("YTD")
    if ytd_start is None:
        ytd_start = pv.index.min()

    portfolio_ytd = portfolio_value[portfolio_value.index >= ytd_start]

    portfolio_ytd = (portfolio_ytd / portfolio_ytd.iloc[0] - 1) * 100

    # Define benchmark tickers and fetch historical prices
    benchmarks = {
        "S&P 500 (^GSPC)": "^GSPC",
        "AOR (Global 60/40)": "AOR",
        "AOK (Conservative 40/60)": "AOK"
    }

    benchmark_prices = {}
    for name, ticker in benchmarks.items():
        hist = fetch_price_history([ticker])
        hist.index = pd.to_datetime(hist.index)
        benchmark_prices[name] = hist[ticker]

    # Slice and compute MTD/YTD for benchmarks
    benchmark_returns = {}
    for name, series in benchmark_prices.items():
        # MTD (MUST USE twr_anchor)
        mtd_series = series[series.index >= twr_anchor]
        mtd_cum = (mtd_series / mtd_series.iloc[0] - 1) * 100

        # YTD
        ytd_series = series[series.index >= ytd_start]
        ytd_cum = (ytd_series / ytd_series.iloc[0] - 1) * 100
        benchmark_returns[name] = {"MTD": mtd_cum, "YTD": ytd_cum}

    # ---------------- MTD TWR-BASED CHART ----------------
    doc.add_page_break()
    doc.add_heading("Portfolio vs Benchmarks Analysis", level=1)
    doc.add_heading("MTD Cumulative Return — Portfolio vs Benchmarks", level=2)

    # 1) Portfolio TWR cumulative curve (based on pv)
    mtd_start_twr = twr_anchor

    # Correct GIPS-compliant daily TWR (flows at start of day)
    cf_ext_local = load_cashflows_external().copy()
    daily_ret = []

    pv_dates = pv.index
    for i in range(1, len(pv_dates)):
        d0 = pv_dates[i-1]
        d1 = pv_dates[i]

        flow = cf_ext_local.loc[cf_ext_local["date"] == d1, "amount"].sum()
        denom = pv.loc[d0] + flow
        if denom <= 0:
            continue

        R = (pv.loc[d1] - denom) / denom
        daily_ret.append((d1, R))


    # Convert to cumulative TWR curve
    twr_curve = pd.Series(1.0, index=[d for d,_ in daily_ret])
    running = 1.0
    for d,R in daily_ret:
        running *= (1 + R)
        twr_curve.loc[d] = running

    # Unified since-inception anchor for both SI chart and SI benchmark table
    si_start = inception_date


    # Slice only the MTD part
    twr_mtd = twr_curve[twr_curve.index >= mtd_start_twr]
    twr_mtd = (twr_mtd / twr_mtd.iloc[0] - 1) * 100


    # 2) Benchmarks: ALSO anchored to the same start date as portfolio
    benchmark_map = {
        "S&P 500 (^GSPC)": "^GSPC",
        "AOR (Global 60/40)": "AOR",
        "AOK (Conservative 40/60)": "AOK"
    }

    benchmark_curves_mtd = {}

    for name, ticker in benchmark_map.items():
        try:
            hist = fetch_price_history([ticker])
            hist.index = pd.to_datetime(hist.index)
            ser = hist[ticker]

            # Align to portfolio’s actual curve start
            ser = ser[ser.index >= mtd_start_twr]

            if len(ser) > 1:
                ser_norm = (ser / ser.iloc[0] - 1.0) * 100.0
                benchmark_curves_mtd[name] = ser_norm
        except:
            pass

    # ---------------- UPDATED MTD CHART (Option C sizing + matching fonts) ----------------
    fig, ax = plt.subplots(figsize=(12, 6.5))

    # Portfolio curve
    ax.plot(
        twr_mtd.index,
        twr_mtd.values,
        linewidth=2,
        label="Portfolio (TWR-Based)"
    )

    # Benchmarks
    for name, series in benchmark_curves_mtd.items():
        ax.plot(
            series.index,
            series.values,
            linewidth=1.6,
            label=name
        )

    # ======== MATCH BAR CHART FONTS EXACTLY ========
    ax.set_title("")
    ax.set_ylabel("Cumulative Return (%)", fontsize=11)
    ax.tick_params(axis="both", labelsize=10)
    ax.legend(fontsize=10)

    plt.tight_layout()

    # Save and insert
    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6.5))   # <<< bigger but not full-page
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_paragraph(
        "Figure: Portfolio returns use true TWR based on a flow-adjusted, end-of-day PV series (external flows treated at start-of-day). Benchmark returns use simple price return and are rebased to the same start date for comparability.",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    # ---------------- SINCE INCEPTION CHART (TWR vs Benchmarks) ----------------
    fig, ax = plt.subplots(figsize=(12, 6.5))

    import matplotlib.dates as mdates

    # Dynamic locator but limited tick count
    locator = mdates.AutoDateLocator(minticks=6, maxticks=6)
    ax.xaxis.set_major_locator(locator)

    # Force full date format
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))

    # Portfolio since-inception TWR curve (twr_curve is already cumulative)
    twr_si_curve = twr_curve.copy()
    twr_si_pct = (twr_si_curve - 1.0) * 100.0


    ax.plot(
        twr_si_pct.index,
        twr_si_pct.values,
        label="Portfolio (TWR Since Inception)",
        linewidth=2,
    )

    # Benchmarks since inception — price return rebased to si_start
    benchmark_map = {
        "S&P 500 (^GSPC)": "^GSPC",
        "AOR (Global 60/40)": "AOR",
        "AOK (Conservative 40/60)": "AOK",
    }

    benchmark_curves_si = {}

    for name, ticker in benchmark_map.items():
        try:
            hist = fetch_price_history([ticker])
            hist.index = pd.to_datetime(hist.index)
            ser = hist[ticker].dropna()

            # Align to portfolio SI start
            ser = ser[ser.index >= si_start]

            if len(ser) > 1:
                ser_norm = (ser / ser.iloc[0] - 1.0) * 100.0
                benchmark_curves_si[name] = ser_norm
        except Exception:
            continue

    for name, series in benchmark_curves_si.items():
        ax.plot(series.index, series.values, label=name, linewidth=1.6)

    # ---------- FONT FIXES (MATCH MTD + BAR CHART) ----------
    ax.set_title("")
    ax.set_ylabel("Cumulative Return (%)", fontsize=11)
    ax.tick_params(axis="both", labelsize=10)
    ax.legend(fontsize=10)

    plt.tight_layout()
    fig.autofmt_xdate()

    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)

    doc.add_heading("Since Inception Cumulative Return — Portfolio vs Benchmarks", level=2)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6.5))   # MATCH MTD WIDTH EXACTLY
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Figure: Portfolio curve uses true since-inception TWR based on the flow-adjusted PV series. "
        "Benchmarks use simple price returns rebased to the same since-inception start date.",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # ---------------------------------------------------------------
    # PERFORMANCE VS BENCHMARKS (ALL HORIZONS — TWR CONSISTENT)
    # ---------------------------------------------------------------
    doc.add_heading("Performance vs Benchmarks (All Horizons)", level=2)

    # Horizons to show (must match your snapshot / twr_df + SI)
    horizons_all = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "SI"]

    # Portfolio TWR (numeric, decimals) from twr_raw
    port_ret_num = {row["Horizon"]: row["Return"] for _, row in twr_raw.iterrows()}
    port_ret_str = {row["Horizon"]: row["Return"] for _, row in twr_df.iterrows()}

    # Inject SI (since inception) into both maps
    if not pd.isna(twr_si):
        port_ret_num["SI"] = twr_si
        # use non-annualized SI TWR here for the table
        port_ret_str["SI"] = fmt_pct_clean(twr_si)

    # Benchmarks and their tickers (keys used for lookups below)
    bm_defs = {
        "S&P 500 %": "^GSPC",
        "Global 60/40 %": "AOR",
        "Conservative 40/60 %": "AOK",
    }

    # Cache benchmark price history once
    bm_prices = {}
    for col_label, ticker in bm_defs.items():
        try:
            hist = fetch_price_history([ticker])
            ser = hist[ticker].dropna()
        except Exception:
            ser = pd.Series(dtype=float)
        bm_prices[col_label] = ser

    # Benchmark return as a decimal, rebased to same start as portfolio
    def compute_bm_ret_decimal(col_label: str, horizon: str) -> float:
        ser = bm_prices[col_label]
        if ser.empty:
            return np.nan

        if horizon == "SI":
            start = si_start  # same since-inception start used in the SI chart
        else:
            start = get_horizon_start(horizon)

        if start is None:
            return np.nan

        ser_h = ser[ser.index >= start]
        if len(ser_h) < 2:
            return np.nan

        return ser_h.iloc[-1] / ser_h.iloc[0] - 1.0  # decimal

    # Build table rows with excess return columns
    combined_rows = []
    for h in horizons_all:
        port_str = port_ret_str.get(h, "N/A")
        port_dec = port_ret_num.get(h, np.nan)

        sp_dec   = compute_bm_ret_decimal("S&P 500 %", h)
        g6040_dec = compute_bm_ret_decimal("Global 60/40 %", h)
        c4060_dec = compute_bm_ret_decimal("Conservative 40/60 %", h)

        # Benchmark display strings
        sp_str    = fmt_pct_clean(sp_dec)
        g6040_str = fmt_pct_clean(g6040_dec)
        c4060_str = fmt_pct_clean(c4060_dec)

        # Excess = portfolio − benchmark (all decimals)
        if pd.isna(port_dec) or pd.isna(sp_dec):
            sp_excess_str = "N/A"
        else:
            sp_excess_str = fmt_pct_clean(port_dec - sp_dec)

        if pd.isna(port_dec) or pd.isna(g6040_dec):
            g6040_excess_str = "N/A"
        else:
            g6040_excess_str = fmt_pct_clean(port_dec - g6040_dec)

        if pd.isna(port_dec) or pd.isna(c4060_dec):
            c4060_excess_str = "N/A"
        else:
            c4060_excess_str = fmt_pct_clean(port_dec - c4060_dec)

        combined_rows.append([
            h,
            port_str,
            sp_str,
            sp_excess_str,
            g6040_str,
            g6040_excess_str,
            c4060_str,
            c4060_excess_str,
        ])

    # Shorten last header label to stop wrapping
    headers = [
        "Horizon",
        "Portfolio %",
        "S&P 500 %",
        "Excess",
        "Global 60/40 %",
        "Excess",
        "Cons 40/60 %",
        "Excess",
    ]

    add_table(
        doc,
        headers,
        combined_rows,
        right_align=[1, 2, 3, 4, 5, 6, 7]
    )
    
    # ---- FIX WIDTHS ONLY FOR THIS BENCHMARKS TABLE ----
    table = doc.tables[-1]          # get last added table
    table.autofit = False
    table.allow_autofit = False

    col_widths = [
        Inches(0.9),  # Horizon
        Inches(0.9),  # Portfolio %
        Inches(0.9),  # S&P 500 %
        Inches(0.9),  # Excess vs S&P
        Inches(1.15),  # Global 60/40 %
        Inches(0.9),  # Excess vs 60/40
        Inches(1.1),  # Cons 40/60 %
        Inches(0.9),  # Excess vs 40/60
    ]

    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w


    doc.add_paragraph(
        "Portfolio % uses flow-adjusted TWR. Benchmark % uses simple price returns "
        "rebased to the same start date for each horizon. Excess = Portfolio % minus benchmark % for the same horizon.",
        style="Normal"
    )

    doc.add_page_break()



    # =============================================================
    # NEW: SINCE-INCEPTION (SI) RETURNS FOR TICKERS & ASSET CLASSES
    # =============================================================

    # We want a true SI return per ticker (since first trade),
    # and a value-weighted rollup for each asset class.
    si_return_map = {}

    tx_all = load_transactions_raw().copy()
    as_of_port = pv.index.max()

    # Use the same price history we already fetched for tickers
    # (non-CASH only; CASH gets 0%).
    for t in sec_only["ticker"].unique():
        if t == "CASH":
            si_return_map[t] = 0.0
            continue

        if t not in prices.columns:
            # No price history → treat as 0% to avoid N/A spam
            si_return_map[t] = 0.0
            continue

        price_series = prices[t].dropna()
        if price_series.empty:
            si_return_map[t] = 0.0
            continue

        tx_t = tx_all[tx_all["ticker"] == t].copy()
        if tx_t.empty:
            # No transactions recorded → fallback to 0%
            si_return_map[t] = 0.0
            continue

        tx_t = tx_t.sort_values("date")
        first_trade = tx_t["date"].min()

        # End date = min(portfolio PV max date, price series max date)
        as_of_price = price_series.index.max()
        end = min(as_of_port, as_of_price)

        if end <= first_trade:
            si_return_map[t] = 0.0
            continue

        # Run MD from first trade date
        try:
            si_ret = modified_dietz_for_ticker_window(
                t,
                price_series,
                tx_t,
                first_trade,
                end,
            )
        except Exception:
            si_ret = np.nan

        # If MD blows up to NaN, leave it as NaN so it shows as "N/A"
        if pd.isna(si_ret):
            si_return_map[t] = np.nan
        else:
            si_return_map[t] = float(si_ret)


    # Attach SI returns to security-level table as a DECIMAL,
    # then we will format for display in the horizon table only.
    if not sec_full.empty:
        sec_full["SI"] = sec_full["ticker"].map(lambda t: si_return_map.get(t, 0.0))

    # Now roll these up to asset classes, value-weighted by current MV.
    if not class_full.empty:
        # Merge current market value into a helper frame
        sec_mv = sec_only[["ticker", "asset_class", "value"]].copy()
        sec_mv["SI"] = sec_mv["ticker"].map(lambda t: si_return_map.get(t, 0.0))

        si_by_class = {}
        for ac, grp in sec_mv.groupby("asset_class"):
            grp = grp.dropna(subset=["value"])
            if grp.empty:
                si_by_class[ac] = 0.0
                continue

            # Only tickers with non-null SI
            sub = grp.dropna(subset=["SI"])
            if sub.empty:
                si_by_class[ac] = 0.0
                continue

            w = sub["value"] / sub["value"].sum()
            si_by_class[ac] = float((w * sub["SI"]).sum())

        # Add SI as decimal to class_full
        class_full["SI"] = class_full["asset_class"].map(
            lambda ac: si_by_class.get(ac, 0.0)
        )

    # Format ticker-level SI for display in the horizon table
    if not sec_full.empty and "SI" in sec_full.columns:
        sec_full["SI"] = sec_full["SI"].apply(fmt_pct_clean)


    # =============================================================
    # MULTI-HORIZON RETURNS BY ASSET CLASS & TICKER (MODIFIED DIETZ)
    # =============================================================
    doc.add_heading("Horizon Performance Tables", level=1)

    doc.add_heading("Returns Grouped by Asset Class (Modified Dietz)", level=2)

    horizon_cols = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "SI"]

    if class_full.empty and sec_full.empty:
        doc.add_paragraph("No security-level cashflow data available.")
    else:
        # Make a formatted view of class_full so we don't double-multiply anything
        class_view = class_full.copy()
        for col in horizon_cols:
            if col in class_view.columns and col != "SI":
                # Non-SI horizons: standard formatting
                class_view[col] = class_view[col].apply(fmt_pct_clean)

        # SI is a decimal; format it once here so asset-class SI shows as %.
        if "SI" in class_view.columns:
            class_view["SI"] = class_view["SI"].apply(fmt_pct_clean)


        # Order asset classes based on class_full
        asset_classes = (
            class_view["asset_class"]
            .dropna()
            .drop_duplicates()
            .tolist()
        )

        rows = []

        for ac in asset_classes:
            # ----- Asset class row -----
            class_row = class_view[class_view["asset_class"] == ac]
            if not class_row.empty:
                r = class_row.iloc[0]
                row_vals = [asset_class_map.get(ac, ac)]
                for h in horizon_cols:
                    row_vals.append(safe(r.get(h)))
                rows.append(row_vals)

            # ----- Ticker rows under that class -----
            tickers_in_ac = (
                sec_full[sec_full["asset_class"] == ac]
                .sort_values("ticker")["ticker"]
                .tolist()
            )

            for t in tickers_in_ac:
                tr = sec_full[sec_full["ticker"] == t].iloc[0]
                t_vals = [f"  {t}"]
                for h in horizon_cols:
                    t_vals.append(safe(tr.get(h)))
                rows.append(t_vals)

        add_table(
            doc,
            ["Asset Class / Ticker"] + horizon_cols,
            rows,
            right_align=list(range(1, 1 + len(horizon_cols)))
        )
        
        # ---- FIX WIDTHS ONLY FOR THIS HORIZON RETURNS TABLE ----
        table = doc.tables[-1]  # get last added table
        table.autofit = False
        table.allow_autofit = False

        col_widths = [
            Inches(1.4),  # Asset Class / Ticker
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8),
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)
        ]

        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w


        doc.add_paragraph(
            "Returns above are computed using a per-holding and per-asset-class Modified Dietz methodology for each horizon.",
            style="Normal"
        )

    # =============================================================
    # MULTI-HORIZON P/L BY ASSET CLASS & TICKER (ECONOMIC, MD-CONSISTENT)
    # =============================================================

    doc.add_heading("P/L Grouped by Asset Class", level=2)

    horizons_pl = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "SI"]

    if class_full.empty and sec_full.empty:
        doc.add_paragraph("No security-level cashflow data available.")
    else:

        def compute_asset_class_pl(asset_class: str, h: str):
            """
            Economic P/L for an asset class over horizon h.

            Computed as the sum of ticker-level economic P/L for all tickers
            in that asset class, where each ticker P/L is already horizon-
            aligned to the same get_horizon_start() logic.

            This is additive and therefore fully consistent with the ticker-
            level Modified Dietz construction.
            """
            # Use sec_only for the universe of tickers and asset_class mapping
            tickers_in_ac = (
                sec_only[sec_only["asset_class"] == asset_class]["ticker"]
                .dropna()
                .unique()
            )

            total_pl = 0.0
            seen_any = False

            for t in tickers_in_ac:
                pl_str = compute_ticker_pl(t, h)
                if pl_str in (None, "N/A"):
                    continue
                try:
                    # parse formatted "$x,xxx.xx" back to float
                    val = float(pl_str.replace("$", "").replace(",", ""))
                except Exception:
                    continue
                total_pl += val
                seen_any = True

            if not seen_any:
                return "N/A"
            return fmt_dollar_clean(total_pl)

        rows_pl = []

        # Use the same asset class ordering as the returns table
        asset_classes_pl = (
            class_full["asset_class"]
            .dropna()
            .drop_duplicates()
            .tolist()
        )

        for ac in asset_classes_pl:
            # ----- Asset class row -----
            row_vals = [asset_class_map.get(ac, ac)]
            for h in horizons_pl:
                row_vals.append(compute_asset_class_pl(ac, h))
            rows_pl.append(row_vals)

            # ----- Ticker rows under that class -----
            tickers_in_ac = (
                sec_only[sec_only["asset_class"] == ac]
                .sort_values("ticker")["ticker"]
                .tolist()
            )

            for t in tickers_in_ac:
                t_vals = [f"  {t}"]
                for h in horizons_pl:
                    t_vals.append(compute_ticker_pl(t, h))
                rows_pl.append(t_vals)

        add_table(
            doc,
            ["Asset Class / Ticker"] + horizons_pl,
            rows_pl,
            right_align=list(range(1, 1 + len(horizons_pl)))
        )
        
        # ---- FIX WIDTHS ONLY FOR THIS HORIZON P/L TABLE ----
        table = doc.tables[-1]  # get last added table
        table.autofit = False
        table.allow_autofit = False

        col_widths = [
            Inches(1.4),  # Asset Class / Ticker
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8),
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)
        ]

        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w


        doc.add_paragraph(
            "P/L above is economic P/L by horizon (MV_end − MV_start − net internal flows), "
            "rolled up additively from tickers to asset classes using the same horizon definitions as the Modified Dietz returns.",
            style="Normal"
        )


    # ---------------------------------------------------------------
    # DETAILED CASH FLOW REPORT (Since Inception)
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Detailed Cash Flow Report", level=1)

    # ===============================================================
    # EXTERNAL CASH FLOWS (Deposits & Withdrawals)
    # ===============================================================
    doc.add_heading("External Cash Flows (Since Inception)", level=2)

    cf_all = load_cashflows_external().copy()
    cf_all = cf_all.sort_values("date")

    if cf_all.empty:
        doc.add_paragraph("No external cash flows recorded.")
    else:
        # Total deposits, withdrawals, net
        total_deposits = cf_all.loc[cf_all["amount"] > 0, "amount"].sum()
        total_withdrawals = cf_all.loc[cf_all["amount"] < 0, "amount"].sum()
        net_ext = cf_all["amount"].sum()

        most_recent = cf_all["date"].max().strftime("%Y-%m-%d") if not cf_all.empty else "N/A"

        ext_rows = [
            ["Total Deposits", fmt_dollar_clean(total_deposits)],
            ["Total Withdrawals", fmt_dollar_clean(total_withdrawals)],
            ["Net External Flow", fmt_dollar_clean(net_ext)],
            ["Most Recent External Flow", most_recent],
        ]

        add_table(
            doc,
            ["Metric", "Value"],
            ext_rows,
            right_align=[1]
        )

    # Small space before internal section
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # ===============================================================
    # INTERNAL TRADING SUMMARY (Buys, Sells, Net)
    # ===============================================================
    doc.add_heading("Internal Trading Summary (Since Inception)", level=2)

    tx_all = load_transactions_raw().copy()
    tx_all = tx_all.sort_values("date")

    if tx_all.empty:
        doc.add_paragraph("No internal trading activity available.")
    else:
        # Aggregate buys (negative amounts) and sells (positive amounts) per ticker
        summary = (
            tx_all.groupby("ticker")["amount"]
            .agg([
                ("buys", lambda s: s[s < 0].sum()),
                ("sells", lambda s: s[s > 0].sum())
            ])
            .reset_index()
        )

        summary["net"] = summary["buys"] + summary["sells"]

        # Format rows for display
        rows = []
        for _, r in summary.iterrows():
            rows.append([
                r["ticker"],
                fmt_dollar_clean(r["buys"]),
                fmt_dollar_clean(r["sells"]),
                fmt_dollar_clean(r["net"])
            ])

        add_table(
            doc,
            ["Ticker", "Total Buys (Cash Out)", "Total Sells (Cash In)", "Net (Cash Flow)"],
            rows,
            right_align=[1, 2, 3]
        )

    # ---------------------------------------------------------------
    # INTERNAL TRADING FLOWS — STACKED BY ASSET CLASS (SINCE INCEPTION)
    # ---------------------------------------------------------------
    doc.add_heading("Internal Trading Flows by Asset Class (Since Inception)", level=2)

    tx_raw = load_transactions_raw().copy()

    if tx_raw.empty:
        doc.add_paragraph("No internal buy/sell activity recorded.")
    else:
        # Map tickers → asset class
        asset_map = (
            holdings_df[["ticker", "asset_class"]]
            .set_index("ticker")["asset_class"]
            .to_dict()
        )
        tx_raw["asset_class"] = tx_raw["ticker"].map(asset_map).fillna("Unknown")

        # Net flows per class
        net_by_class = (
            tx_raw.groupby("asset_class")["amount"]
            .sum()
            .sort_values(ascending=True)
        )

        # Build stacked bar (fixed)
        fig, ax = plt.subplots(figsize=(9, 5.5))


        left_edge = 0
        labels = []

        for cls, val in net_by_class.items():
            ax.barh(
                ["Net Flows"],
                val,
                left=left_edge,
                label=f"{cls} (${val:,.0f})"
            )
            left_edge += val

        # FORCE FULL RANGE TO SHOW ENTIRE BAR
        max_right = max(left_edge, 0)
        min_left = min(0, left_edge)
        ax.set_xlim(min_left * 1.05, max_right * 1.05)

        ax.set_title("")

        ax.set_xlabel("Net Cash Flow ($)")
        ax.grid(axis="x", linestyle="--", alpha=0.4)

        # Make legend readable
        ax.legend(
            loc="upper center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=3,
            fontsize=9.5
        )

        plt.tight_layout()

        buf = BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(buf, width=Inches(6))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # PV MOUNTAIN CHART (Normalized to % Return Since Inception)
    # ---------------------------------------------------------------
    doc.add_heading("Portfolio Value Analysis", level=1)

    doc.add_heading("Portfolio Value — Mountain Chart (Since-Inception % Return)", level=2)

    # Build daily PV and forward-fill
    pv_daily = pv.sort_index().reindex(
        pd.date_range(pv.index.min(), pv.index.max(), freq="D")
    ).ffill()

    # Convert to % return since inception
    pv0 = pv_daily.iloc[0]
    pv_ret = (pv_daily / pv0 - 1.0) * 100.0

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(pv_ret.index, pv_ret.values, linewidth=2)
    ax.fill_between(pv_ret.index, pv_ret.values, alpha=0.25)

    ax.set_title("")
    ax.set_ylabel("Return (%)")
    ax.grid(alpha=0.3)

    # Limit x-axis to 6 ticks
    xticks = np.linspace(0, len(pv_ret.index) - 1, 6, dtype=int)
    ax.set_xticks(pv_ret.index[xticks])
    fig.autofmt_xdate()

    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Figure: Portfolio value normalized to % return since inception. "
        "Shows true performance shape even when dollar PV is flat.",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # DAILY ΔPV ATTRIBUTION — EXTERNAL FLOWS vs MARKET MOVE
    # ---------------------------------------------------------------
    doc.add_heading("Daily ΔPV Attribution", level=2)

    # Build daily PV (fill missing days)
    pv_trading = pv.sort_index().copy()
    pv_daily = pv_trading.reindex(
        pd.date_range(pv_trading.index.min(), pv_trading.index.max(), freq="D")
    ).ffill()

    # Daily ΔPV
    dpv = pv_daily.diff().fillna(0.0)

    # ---------------- EXTERNAL FLOWS ----------------
    cf_all = load_cashflows_external().copy()
    if not cf_all.empty:
        cf_all = cf_all.sort_values("date")
        ext_series = (
            cf_all.groupby("date")["amount"]
            .sum()
            .reindex(pv_daily.index, fill_value=0.0)
        )
    else:
        ext_series = pd.Series(0.0, index=pv_daily.index)

    # ---------------- INTERNAL FLOWS (Buys/Sells) ----------------
    tx_all = load_transactions_raw().copy()
    if not tx_all.empty:
        tx_all = tx_all.sort_values("date")
        internal_series = (
            tx_all.groupby("date")["amount"]
            .sum()
            .reindex(pv_daily.index, fill_value=0.0)
        )
    else:
        internal_series = pd.Series(0.0, index=pv_daily.index)
        

    # ---------------- MARKET EFFECT ----------------
    # Residual after removing BOTH external + internal flows
    mkt_series = dpv - ext_series - internal_series


    # Last 30 days window
    window_days = 30
    end_date = pv_daily.index.max()
    start_date = max(end_date - pd.Timedelta(days=window_days - 1), pv_daily.index.min())

    mask = (pv_daily.index >= start_date) & (pv_daily.index <= end_date)
    dates_win = pv_daily.index[mask]
    dpv_win = dpv[mask]
    ext_win = ext_series[mask]
    internal_win = internal_series[mask]
    mkt_win = mkt_series[mask]

    # Cumulative ΔPV (rebased to 0)
    cum_total = (ext_series + internal_series + mkt_series).cumsum()
    cum_win = cum_total[mask] - cum_total[mask].iloc[0]

    import matplotlib.dates as mdates

    fig, ax1 = plt.subplots(figsize=(9, 4.5))

    # External flows layer
    ax1.bar(
        dates_win,
        ext_win,
        label="External Flows",
        width=0.9,
    )

    # Internal flows layer (stacked on external)
    ax1.bar(
        dates_win,
        internal_win,
        bottom=ext_win,
        label="Internal Flows",
        width=0.9,
    )

    # Market effects layer (stacked on external + internal)
    ax1.bar(
        dates_win,
        mkt_win,
        bottom=ext_win + internal_win,
        label="Market Effects",
        width=0.9,
    )


    ax1.set_ylabel("Δ Portfolio Value ($)")
    ax1.set_title("")

    # -----------------------------------------------------------
    # X AXIS — EXACT YYYY-MM-DD AND EXACTLY 6 TICKS
    # -----------------------------------------------------------
    num_xticks = 6
    tick_positions = np.linspace(
        mdates.date2num(dates_win[0]),
        mdates.date2num(dates_win[-1]),
        num_xticks
    )

    ax1.set_xticks(tick_positions)
    ax1.set_xticklabels(
        [mdates.num2date(t).strftime("%Y-%m-%d") for t in tick_positions],
        rotation=35,
        ha="right"
    )

    # -----------------------------------------------------------
    # Y AXIS — AUTO ZERO-CENTER FOR DELTA PV
    # -----------------------------------------------------------
    max_abs = max(abs(dpv_win.min()), abs(dpv_win.max()))
    ax1.set_ylim(-max_abs * 1.1, max_abs * 1.1)

    # -----------------------------------------------------------
    # RIGHT AXIS — CUMULATIVE ΔPV
    # -----------------------------------------------------------
    ax2 = ax1.twinx()
    ax2.plot(
        dates_win,
        cum_win,
        linestyle="--",
        marker="o",
        linewidth=1.8,
        label="Cumulative ΔPV",
    )
    ax2.set_ylabel("Cumulative ΔPV ($)")
    ax2.yaxis.set_major_locator(plt.MaxNLocator(6))

    # Left y-axis ticks
    ax1.yaxis.set_major_locator(plt.MaxNLocator(6))

    # Combine legends
    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(
        handles1 + handles2,
        labels1 + labels2,
        loc="upper left",
        fontsize=8,
    )

    plt.tight_layout()

    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph(
        "Figure: Daily decomposition of portfolio value changes into external flows "
        "and market movements. Stacked bars sum to total ΔPV; the dashed line shows "
        "cumulative ΔPV over the same window. Y-axis is auto-centered on zero.",
        style="Normal",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---------------------------------------------------------------
    # LONG-TERM PROJECTION SCENARIOS (20 YEARS)
    # ---------------------------------------------------------------
    doc.add_heading("20-Year Projection Scenarios", level=1)

    # Load contribution amount from config
    monthly_contrib = TARGET_MONTHLY_CONTRIBUTION

    # Initial portfolio value = current PV
    initial_value = float(pv.iloc[-1])

    # Projection assumptions
    years = [1, 5, 10, 15, 20]
    rates = [0.05, 0.07, 0.09]

    # Helper: future value of lump sum
    def fv_lump(pv0, r, yr):
        return pv0 * ((1 + r) ** yr)

    # Helper: future value of monthly contributions
    def fv_contrib(c, r, yr):
        monthly_r = r / 12.0
        n = yr * 12
        if monthly_r == 0:
            return c * n
        return c * (( (1 + monthly_r) ** n - 1 ) / monthly_r)

    # Build projection table rows
    proj_rows = []
    for yr in years:
        row = [yr]
        # Lump-only
        for r in rates:
            row.append(f"${fv_lump(initial_value, r, yr):,.0f}")
        # Lump + monthly contributions
        for r in rates:
            fv_total = fv_lump(initial_value, r, yr) + fv_contrib(monthly_contrib, r, yr)
            row.append(f"${fv_total:,.0f}")
        proj_rows.append(row)

    # Build table
    headers = [
        "Year",
        "5% (Lump)",
        "7% (Lump)",
        "9% (Lump)",
        f"5% (+${monthly_contrib:,.0f}/mo)",
        f"7% (+${monthly_contrib:,.0f}/mo)",
        f"9% (+${monthly_contrib:,.0f}/mo)"
    ]

    add_table(
        doc,
        headers,
        proj_rows,
        right_align=[1,2,3,4,5,6]
    )

    # Add spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)


    # ---------------------------------------------------------------
    # LONG-TERM PROJECTION CHART
    # ---------------------------------------------------------------

    # Year-by-year curves for full plotting
    horizon_years = range(0, 21)

    proj_lump = {r: [] for r in rates}
    proj_contrib = {r: [] for r in rates}

    for yr in horizon_years:
        for r in rates:
            proj_lump[r].append(fv_lump(initial_value, r, yr))
            proj_contrib[r].append(fv_lump(initial_value, r, yr) + fv_contrib(monthly_contrib, r, yr))

    fig, ax = plt.subplots(figsize=(8, 5))

    # Plot lump sum lines (solid)
    ax.plot(horizon_years, proj_lump[0.05], label="5% Lump Sum", linewidth=2)
    ax.plot(horizon_years, proj_lump[0.07], label="7% Lump Sum", linewidth=2)
    ax.plot(horizon_years, proj_lump[0.09], label="9% Lump Sum", linewidth=2)

    # Plot contrib lines (dotted)
    ax.plot(horizon_years, proj_contrib[0.05], linestyle="--", label=f"5% + ${monthly_contrib:,.0f}/mo", linewidth=2)
    ax.plot(horizon_years, proj_contrib[0.07], linestyle="--", label=f"7% + ${monthly_contrib:,.0f}/mo", linewidth=2)
    ax.plot(horizon_years, proj_contrib[0.09], linestyle="--", label=f"9% + ${monthly_contrib:,.0f}/mo", linewidth=2)

    ax.set_xlabel("Years")
    ax.set_ylabel("Portfolio Value ($)")

    # Fix Y-axis to show normal dollars, not scientific notation
    ax.ticklabel_format(style='plain', axis='y')
    ax.get_yaxis().set_major_formatter(plt.FuncFormatter(lambda x, pos: f"${x:,.0f}"))

    ax.set_title("")
    ax.legend(fontsize=8)
    plt.tight_layout()

    # Save to memory and insert into doc
    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph("Figure: Long-term projections.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # RISK & VOLATILITY ANALYSIS (APPENDIX)
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Risk & Volatility Analysis", level=1)

    # ===========================
    # Expected Volatility By Asset Class (Bar Chart)
    # ===========================
    doc.add_heading("Expected Volatility by Asset Class", level=2)

    vol_data = {
        "Digital Assets": 75,
        "Gold / Precious Metals": 14,
        "International Equity": 17.5,
        "US Bonds": 6,
        "US Growth": 19,
        "US Large Cap": 15,
        "US Small Cap": 24
    }

    classes = list(vol_data.keys())
    vol_vals = list(vol_data.values())

    fig, ax = plt.subplots(figsize=(8,5))
    bars = ax.bar(classes, vol_vals)

    ax.set_title("")
    ax.set_ylabel("Std Dev (%)")

    plt.xticks(rotation=25, ha="right")
    plt.tight_layout()

    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph("Figure: Approximate volatility estimate.", style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===========================
    # Risk vs Expected Return (Scatter)
    # ===========================
    doc.add_heading("Risk vs Expected Return", level=2)

    exp_return_data = {
        "Digital Assets": (75, 12),
        "US Small Cap": (24, 9.0),
        "US Growth": (19, 8.5),
        "International Equity": (17.5, 8.0),
        "US Large Cap": (15, 7.0),
        "Gold / Precious Metals": (14.5, 5.5),
        "US Bonds": (6, 4.5)
    }

    vols = [v[0] for v in exp_return_data.values()]
    rets = [v[1] for v in exp_return_data.values()]
    labels = list(exp_return_data.keys())

    fig, ax = plt.subplots(figsize=(8,5))
    ax.scatter(vols, rets, s=80)

    for i, label in enumerate(labels):
        ax.annotate(label, (vols[i] + 0.5, rets[i] + 0.1), fontsize=9)

    ax.set_xlabel("Volatility (%)")
    ax.set_ylabel("Expected Annual Return (%)")
    ax.set_title("")
    ax.grid(True, linestyle="--", alpha=0.4)

    plt.tight_layout()

    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph("Figure: Trade-off between expected return and volatility.", style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ---------------------------------------------------------------
    # SECTOR ALLOCATION HEATMAP
    # ---------------------------------------------------------------
    doc.add_page_break()

    # Hardcode ETF sector map from config
    sector_exposure = defaultdict(float)

    # Sector normalization to avoid duplicates
    SECTOR_NORMALIZATION = {
        "Comm Services": "Communication Services",
        "Consumer Disc.": "Consumer Discretionary",
        "Information Technology": "Tech",
        "Other": None,  # ignore "Other" buckets
    }

    # Use raw market values and exclude CASH from the denominator so that
    # sector weights are based on invested assets only.
    sector_universe = sec_only[sec_only["ticker"].isin(ETF_SECTOR_MAP.keys())].copy()
    sector_universe = sector_universe[sector_universe["value"] > 0]

    total_invested = sector_universe["value"].sum()

    if total_invested > 0:
        for _, row in sector_universe.iterrows():
            ticker = row["ticker"]
            # weight of this holding as % of invested (non-CASH) assets
            weight_pct = (row["value"] / total_invested) * 100.0

            etf_sectors = ETF_SECTOR_MAP.get(ticker, {})
            for sector, pct in etf_sectors.items():
                norm_sector = SECTOR_NORMALIZATION.get(sector, sector)
                if norm_sector is None:
                    continue
                sector_exposure[norm_sector] += weight_pct * pct / 100.0


    # Convert to sorted DataFrame
    sector_df = pd.DataFrame(
        list(sector_exposure.items()),
        columns=["Sector", "Exposure"]
    ).sort_values("Exposure", ascending=True)

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.barh(sector_df["Sector"], sector_df["Exposure"])
    for i, v in enumerate(sector_df["Exposure"]):
        ax.text(v + 0.3, i, f"{v:.1f}%", va="center", fontsize=9)

    ax.set_xlabel("Portfolio Exposure (%)")
    ax.set_title("")
    plt.tight_layout()

    img_stream = BytesIO()
    plt.savefig(img_stream, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    img_stream.seek(0)

    doc.add_heading("Sector Allocation", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure: Approximate sector exposure for portfolio ETFs.", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =============================================================
    # METHODOLOGY APPENDIX (CLEAN ONE-PAGER)
    # =============================================================
    doc.add_page_break()
    doc.add_heading("Appendix — Methodology Summary", level=1)

    # Subtle intro line
    p = doc.add_paragraph()
    r = p.add_run("This page summarizes how the figures in this report are calculated.")
    r.italic = True

    # --- Data & Portfolio Value ---
    p = doc.add_paragraph()
    r = p.add_run("Data and portfolio value. ")
    r.bold = True
    p.add_run(
        "Holdings, transactions, and cash flows are taken from the source files and combined with daily market prices "
        "to build a consistent portfolio value series. External cash flows (deposits and withdrawals) are separated "
        "from internal trading activity (buys and sells) so that performance can be evaluated independently of funding decisions."
    )

    # --- Portfolio Time-Weighted Returns (TWR) ---
    p = doc.add_paragraph()
    r = p.add_run("Portfolio time-weighted returns (TWR). ")
    r.bold = True
    p.add_run(
        "Portfolio performance is reported using a flow-adjusted, time-weighted methodology. "
        "External cash flows are treated as occurring at the start of the day, and the history is broken into segments "
        "between flows. Returns for these segments are chained to produce cumulative TWR over each horizon, consistent "
        "with institutional and GIPS-style practice."
    )

    # --- Security & Asset-Class Returns (Modified Dietz) ---
    p = doc.add_paragraph()
    r = p.add_run("Security and asset-class returns (Modified Dietz). ")
    r.bold = True
    p.add_run(
        "For individual securities and asset classes, the report uses a Modified Dietz money-weighted approach. "
        "All ticker-level cash flows (buys and sells) and prices over the horizon are included, and returns are "
        "computed relative to an average invested capital base. Asset-class returns are formed as value-weighted "
        "averages of the underlying security-level returns."
    )

    # --- Profit and Loss (P/L) ---
    p = doc.add_paragraph()
    r = p.add_run("Profit and loss (P/L). ")
    r.bold = True
    p.add_run(
        "Economic P/L is defined consistently throughout the report as: MV_end minus MV_start minus net cash flows "
        "over the period. At the portfolio level, cash flows are external only (contributions and withdrawals). "
        "At the security and asset-class levels, P/L is based on internal trading flows and aligned to the same "
        "time horizons used for the return calculations."
    )

    # --- Benchmarks & Comparisons ---
    p = doc.add_paragraph()
    r = p.add_run("Benchmarks and comparisons. ")
    r.bold = True
    p.add_run(
        "Benchmark figures (such as broad equity or balanced indices) are shown as simple price returns. "
        "They are rebased to the same start dates as the portfolio for each chart or table so that relative "
        "performance reflects differences in allocation and market exposure rather than timing."
    )

    # --- Data Quality & Limitations ---
    p = doc.add_paragraph()
    r = p.add_run("Data quality and limitations. ")
    r.bold = True
    p.add_run(
        "All results depend on the completeness and accuracy of the underlying holdings, transaction, cash flow, "
        "and price data. Minor differences from broker statements may arise from rounding, pricing gaps, or "
        "timing conventions, but these should not be material to the overall conclusions. The methodology is "
        "designed to be transparent, repeatable, and suitable for executive-level review."
    )

    # Final short design principles list to make it "pop"
    doc.add_paragraph("Core design principles:", style="Normal")
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Separate funding decisions (flows) from investment performance.")
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Use true time-weighted returns at the portfolio level.")
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Use money-weighted, cash-flow-aware math for securities and asset classes.")


    # =============================================================
    # SAVE OUTPUT
    # =============================================================

    DOCX_NAME = f"Portfolio_Performance_Report_{timestamp}.docx"
    PDF_NAME  = f"Portfolio_Performance_Report_{timestamp}.pdf"

    doc.save(DOCX_NAME)

    try:
        convert(DOCX_NAME, PDF_NAME)
    except:
        pass

    print("✔ Report generated:")
    print("   -", DOCX_NAME)
    print("   -", PDF_NAME)


if __name__ == "__main__":
    build_report()
