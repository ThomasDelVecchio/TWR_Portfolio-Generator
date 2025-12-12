#!/usr/bin/env python3
import pandas as pd
import numpy as np
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx2pdf import convert
from io import BytesIO
from collections import defaultdict

# Imports from new modules
from config import (
    TARGET_PORTFOLIO_VALUE, 
    TARGET_MONTHLY_CONTRIBUTION, 
    ETF_SECTOR_MAP
)
from data_loader import (
    load_cashflows_external, 
    load_transactions_raw, 
    load_dividends, 
    fetch_price_history,
    load_holdings
)
from financial_math import (
    get_portfolio_horizon_start, 
    modified_dietz_for_ticker_window, 
    fv_lump, 
    fv_contrib
)
from portfolio_engine import (
    run_engine, 
    calculate_horizon_pl, 
    calculate_ticker_pl
)
from report_formatting import (
    fmt_pct_clean, 
    fmt_dollar_clean, 
    safe, 
    add_table
)
from report_charts import (
    plot_ticker_allocation_pie,
    plot_ticker_allocation_bar,
    plot_asset_allocation_pie,
    plot_asset_allocation_bar,
    plot_asset_allocation_history,
    plot_mtd_cumulative_return,
    plot_si_cumulative_return,
    plot_excess_return,
    plot_internal_trading_flows,
    plot_pv_mountain,
    plot_daily_dpv_attribution,
    plot_long_term_projection,
    plot_expected_volatility,
    plot_risk_return,
    plot_sector_allocation
)

# =====================================================================
# MAIN REPORT BUILDER — FULLY CORRECTED
# =====================================================================

def build_report():

    # Run the engine (unchanged math)
    twr_df, sec_full, class_full, pv, twr_si, twr_si_annualized, pl_si = run_engine()
    
    twr_raw = twr_df.copy()  # keep numeric returns for excess-return calc


    cf_ext = load_cashflows_external()

    tx_raw = load_transactions_raw()

    
    # === Inception date (same logic as run_engine) ===
    dates = []
    if not cf_ext.empty:
        dates.append(cf_ext["date"].min())
    if not tx_raw.empty:
        dates.append(tx_raw["date"].min())
    dates.append(pv.index.min())

    inception_date = min(dates)

    
    # External cashflows for proper P/L (deposits/withdrawals only)
    cf_ext = load_cashflows_external()

    # =============================================================
    # CLEAN ALL RETURN COLUMNS
    # =============================================================

    # Fix TWR DF
    twr_df["Return"] = twr_df["Return"].apply(fmt_pct_clean)

    # Fix Security-level DF
    if not sec_full.empty:
        for col in ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "YTD"]:
            if col in sec_full.columns:
                sec_full[col] = sec_full[col].apply(fmt_pct_clean)

    # =============================================================
    # START DOC
    # =============================================================

    doc = Document()

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    base.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # =============================================================
    # COVER PAGE
    # =============================================================
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
    doc.add_paragraph("\n\n\n\n")

    title = doc.add_paragraph("Portfolio Performance Report")
    title.runs[0].font.size = Pt(26)
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Time-Weighted & Money-Weighted Performance (Automated)")
    created_line = doc.add_paragraph(f"Created for Tom Short — {timestamp}")
    created_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    created_line.runs[0].font.size = Pt(12)
    created_line.runs[0].italic = True
    created_line.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(70, 130, 180)

    doc.add_page_break()

    # =============================================================
    # PORTFOLIO SNAPSHOT (VERTICAL FORMAT + REAL P/L)
    # =============================================================

    doc.add_heading("Portfolio Snapshot", level=1)
    doc.add_paragraph("Summary of time-weighted returns and P/L.")

    # Extract Return % values
    snap_map = {}
    for _, row in twr_df.iterrows():
        snap_map[row["Horizon"]] = row["Return"]

    # Horizons in EXACT order
    horizons = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y"]

    # Build vertical snapshot rows
    rows = []
    for h in horizons:
        ret = snap_map.get(h, "N/A")
        pl_val = calculate_horizon_pl(pv, inception_date, cf_ext, h)
        pl = fmt_dollar_clean(pl_val) if pl_val is not None else "N/A"
        if ret == "N/A": pl = "N/A" # Consistency
        rows.append([h, ret, pl])
        
    # ===== ADD SINCE-INCEPTION ROW =====
    si_pl = pl_si

    # Prefer annualized SI if run_engine computed it (i.e., inception > ~1 year).
    if twr_si_annualized is not None and not pd.isna(twr_si_annualized):
        si_ret_for_display = twr_si_annualized
    else:
        si_ret_for_display = twr_si

    rows.append([
        "Since Inception",
        fmt_pct_clean(si_ret_for_display),
        fmt_dollar_clean(si_pl)
    ])

    add_table(
        doc,
        ["Horizon", "Return %", "P/L ($)"],
        rows,
        right_align=[1, 2]
    )

    # Small spacing before summary table
    doc.add_paragraph().paragraph_format.space_before = Pt(2)

    # Portfolio Summary Table
    total_value = pv.iloc[-1] if isinstance(pv, pd.Series) else pv
    num_holdings = len(sec_full)

    summary_rows = [
        ["Total Value", fmt_dollar_clean(pv.iloc[-1])],
        ["Target Portfolio Value", fmt_dollar_clean(TARGET_PORTFOLIO_VALUE)],
        ["Number of Holdings", num_holdings],
    ]

    add_table(
        doc,
        ["Metric", "Value"],
        summary_rows,
        right_align=[1],
    )

    # PREP FOR PERFORMANCE HIGHLIGHTS — BUILD sec_only + prices

    holdings_df = load_holdings()

    # Merge target_pct into sec_full (include CASH now)
    sec_only = sec_full.copy()
    sec_only = sec_only.merge(
        holdings_df[["ticker", "target_pct"]],
        on="ticker",
        how="left"
    )
    sec_only["target_pct"] = sec_only["target_pct"].fillna(0.0)

    # Shorten asset class names
    asset_class_map = {
        "US Large Cap": "US LC",
        "US Growth": "US Growth",
        "US Small Cap": "US SC",
        "International Equity": "INTL EQTY",
        "Gold / Precious Metals": "GOLD / PM",
        "Digital Assets": "DIGITAL",
        "US Bonds": "US Bonds",
        "CASH": "Cash"
    }
    sec_only["asset_class_short"] = sec_only["asset_class"].map(lambda x: asset_class_map.get(x, x))

    # Fetch live prices for non-cash tickers
    tickers = sec_only[sec_only["ticker"] != "CASH"]["ticker"].tolist()
    prices = fetch_price_history(tickers)
    latest_prices = prices.iloc[-1]
    
    # Pre-load transactions for ticker PL to avoid re-loading inside loop
    tx_all_raw = load_transactions_raw()

    def get_ticker_pl_str(ticker, h):
        as_of = pv.index.max()
        if h == "SI":
            raw_start = None
        else:
            raw_start = get_portfolio_horizon_start(pv, inception_date, h)
        
        pl_val = calculate_ticker_pl(ticker, h, prices, as_of, tx_all_raw, sec_only, raw_start)
        return fmt_dollar_clean(pl_val) if pl_val is not None else "N/A"

    # ---------------------------------------------------------------
    # PERFORMANCE HIGHLIGHTS TABLE
    # ---------------------------------------------------------------
    doc.add_heading("Performance Highlights", level=1)

    perf_rows = []

    if not sec_full.empty:
        # Build numeric version for selecting top/bottom performers
        sec_full_numeric = sec_full.copy()
        for col in ["1M", "1D"]:
            if col in sec_full_numeric.columns:
                sec_full_numeric[col] = (
                    sec_full_numeric[col].astype(str)
                    .str.replace("%", "", regex=False)
                    .replace("N/A", np.nan)
                    .astype(float)
                )

        # ---------------- TOP / BOTTOM 1M ----------------
        if "1M" in sec_full_numeric and not sec_full_numeric["1M"].dropna().empty:
            top_1m = sec_full_numeric.loc[sec_full_numeric["1M"].idxmax()]
            bottom_1m = sec_full_numeric.loc[sec_full_numeric["1M"].idxmin()]

            top_pl_1m = get_ticker_pl_str(top_1m["ticker"], "1M")
            bottom_pl_1m = get_ticker_pl_str(bottom_1m["ticker"], "1M")

            perf_rows.append([
                "Top 1M Performer",
                f"{top_1m['ticker']} ({fmt_pct_clean(top_1m['1M']/100)}, {top_pl_1m})"
            ])
            perf_rows.append([
                "Bottom 1M Performer",
                f"{bottom_1m['ticker']} ({fmt_pct_clean(bottom_1m['1M']/100)}, {bottom_pl_1m})"
            ])
        else:
            perf_rows.append(["Top 1M Performer", "N/A"])
            perf_rows.append(["Bottom 1M Performer", "N/A"])

        # ---------------- BEST / WORST 1D ----------------
        if "1D" in sec_full_numeric and not sec_full_numeric["1D"].dropna().empty:
            best_1d = sec_full_numeric.loc[sec_full_numeric["1D"].idxmax()]
            bottom_1d = sec_full_numeric.loc[sec_full_numeric["1D"].idxmin()]

            best_pl_1d = get_ticker_pl_str(best_1d["ticker"], "1D")
            bottom_pl_1d = get_ticker_pl_str(bottom_1d["ticker"], "1D")

            perf_rows.append([
                "Best 1D Performer",
                f"{best_1d['ticker']} ({fmt_pct_clean(best_1d['1D']/100)}, {best_pl_1d})"
            ])
            perf_rows.append([
                "Bottom 1D Performer",
                f"{bottom_1d['ticker']} ({fmt_pct_clean(bottom_1d['1D']/100)}, {bottom_pl_1d})"
            ])
        else:
            perf_rows.append(["Best 1D Performer", "N/A"])
            perf_rows.append(["Bottom 1D Performer", "N/A"])

    else:
        perf_rows = [
            ["Top 1M Performer", "N/A"],
            ["Bottom 1M Performer", "N/A"],
            ["Best 1D Performer", "N/A"],
            ["Bottom 1D Performer", "N/A"],
        ]

    # Add table
    add_table(
        doc,
        ["Metric", "Value"],
        perf_rows,
        right_align=[1]
    )


    # ---------------------------------------------------------------
    # RISK & DIVERSIFICATION TABLE
    # ---------------------------------------------------------------
    doc.add_heading("Risk & Diversification", level=1)

    # Exclude CASH from weight calculations
    sec_no_cash = sec_full[sec_full["ticker"] != "CASH"].copy()

    # Top 3 holdings % of portfolio
    top3_pct = sec_no_cash.nlargest(3, "weight")["weight"].sum() * 100 if not sec_no_cash.empty else 0

    # Compute total weights per asset class
    asset_class_weights = sec_no_cash.groupby("asset_class")["weight"].sum() * 100

    # Largest asset class by current weight
    largest_class = asset_class_weights.idxmax() if not asset_class_weights.empty else "N/A"
    largest_class_pct = asset_class_weights.max() if not asset_class_weights.empty else 0

    # Compute target percentages from sample holdings
    target_pct_map = holdings_df.groupby("asset_class")["target_pct"].sum().to_dict()

    # Largest over/underweight relative to target
    largest_over = None
    largest_under = None
    max_diff = -np.inf
    min_diff = np.inf

    for ac, wt in asset_class_weights.items():
        target = target_pct_map.get(ac, 0)
        diff = wt - target
        if diff > max_diff:
            max_diff = diff
            largest_over = f"{ac} ({wt:.2f}% vs {target:.2f}%)"
        if diff < min_diff:
            min_diff = diff
            largest_under = f"{ac} ({wt:.2f}% vs {target:.2f}%)"

    risk_rows = [
        ["Top 3 holdings % of portfolio", f"{top3_pct:.2f}%"],
        ["Largest asset class", f"{largest_class} ({largest_class_pct:.2f}%)"],
        ["Largest overweight", largest_over if largest_over else "N/A"],
        ["Largest underweight", largest_under if largest_under else "N/A"],
    ]

    add_table(
        doc,
        ["Metric", "Value"],
        risk_rows,
        right_align=[1]
    )
   

    # ---------------------------------------------------------------
    # FLOWS SUMMARY (Unified External + Internal) — YTD VERSION
    # ---------------------------------------------------------------
    doc.add_heading("Flows Summary (YTD)", level=1)

    as_of = pv.index.max()
    ytd_start = as_of.replace(month=1, day=1)

    # External flows (deposits/withdrawals)
    flows_ext = cf_ext.copy()
    flows_ext = flows_ext[flows_ext["date"] >= ytd_start].sort_values("date")

    if flows_ext.empty:
        ytd_deposits = 0.0
        ytd_withdrawals = 0.0
        net_ytd_ext = 0.0
        most_recent_ext = None
    else:
        ytd_deposits = flows_ext.loc[flows_ext["amount"] > 0, "amount"].sum()
        ytd_withdrawals = flows_ext.loc[flows_ext["amount"] < 0, "amount"].sum()
        net_ytd_ext = flows_ext["amount"].sum()
        most_recent_ext = flows_ext["date"].max()

    # Internal trades (buys/sells)
    tx_ytd = tx_raw.copy()
    tx_ytd = tx_ytd[tx_ytd["date"] >= ytd_start].sort_values("date")

    if tx_ytd.empty:
        ytd_buys = 0.0
        ytd_sells = 0.0
        most_recent_tx = None
    else:
        # buys = negative amounts (cash out)
        ytd_buys = tx_ytd.loc[tx_ytd["amount"] < 0, "amount"].sum()
        ytd_sells = tx_ytd.loc[tx_ytd["amount"] > 0, "amount"].sum()
        most_recent_tx = tx_ytd["date"].max()

    # Income (Dividends)
    div_raw = load_dividends().copy()
    div_raw = div_raw[div_raw["date"] >= ytd_start].sort_values("date")

    if div_raw.empty:
        ytd_income = 0.0
        most_recent_div = None
    else:
        ytd_income = div_raw["amount"].sum()
        most_recent_div = div_raw["date"].max()

    # Net Internal
    net_ytd_internal = ytd_buys + ytd_sells + ytd_income

    # Choose the most recent of ANY flow
    dates_list = []
    if most_recent_ext: dates_list.append(most_recent_ext)
    if most_recent_tx: dates_list.append(most_recent_tx)
    if most_recent_div: dates_list.append(most_recent_div)

    if dates_list:
        most_recent_any = max(dates_list).strftime("%Y-%m-%d")
    else:
        most_recent_any = "N/A"

    flow_rows = [
        ["YTD Net External Flows", fmt_dollar_clean(net_ytd_ext)],
        ["• YTD Deposits", fmt_dollar_clean(ytd_deposits)],
        ["• YTD Withdrawals", fmt_dollar_clean(ytd_withdrawals)],
        ["YTD Net Internal Activity", fmt_dollar_clean(net_ytd_internal)],
        ["• YTD Buys (Cash Out)", fmt_dollar_clean(ytd_buys)],
        ["• YTD Sells (Cash In)", fmt_dollar_clean(ytd_sells)],
        ["• YTD Income (Divs)", fmt_dollar_clean(ytd_income)],
        ["Most Recent Flow", most_recent_any],
    ]

    add_table(
        doc,
        ["Metric", "Value"],
        flow_rows,
        right_align=[1]
    )


    # ---------------------------------------------------------------
    # PORTFOLIO COMPOSITION & STRATEGY TABLE
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Portfolio Composition & Strategy", level=1)
    doc.add_heading("Holdings by Ticker", level=2)

    # Compute prices (CASH = 1)
    sec_only["price"] = sec_only["ticker"].map(lambda t: 1 if t == "CASH" else latest_prices[t])
    sec_only["value"] = sec_only["shares"] * sec_only["price"]
    total_value = sec_only["value"].sum()
    sec_only["allocation"] = sec_only["value"] / total_value * 100

    # Compute "To Contrib" = amount needed to reach target weight, never negative
    sec_only["to_contrib"] = np.maximum(((sec_only["target_pct"] - sec_only["allocation"]) / 100) * total_value, 0)

    # Format columns for display
    # (Create new display columns to avoid breaking numeric cols used below)
    sec_only["allocation_disp"] = sec_only["allocation"].map(lambda x: f"+{x:.2f}%" if x >= 0 else f"{x:.2f}%")
    sec_only["target_pct_disp"] = sec_only["target_pct"].map(lambda x: f"+{x:.2f}%" if x >= 0 else f"{x:.2f}%")
    sec_only["to_contrib_disp"] = sec_only["to_contrib"].map(lambda x: f"${x:,.2f}")

    # Build table rows
    table_rows = []
    for _, row in sec_only.iterrows():
        table_rows.append([
            row["ticker"],
            row.get("asset_class_short", "Unknown"),
            f"{row['shares']:.3f}",
            fmt_dollar_clean(row["price"]),
            fmt_dollar_clean(row["value"]),
            row["allocation_disp"],
            row["target_pct_disp"],
            row["to_contrib_disp"],
        ])

    # Add TOTAL row
    total_value_sum = sec_only["value"].sum()
    total_to_contrib_sum = sec_only["to_contrib"].sum()
    table_rows.append([
        "TOTAL",
        "",
        "",
        "",
        fmt_dollar_clean(total_value_sum),
        "",
        "",
        f"${total_to_contrib_sum:,.2f}"
    ])

    # Add table
    add_table(
        doc,
        ["Ticker", "Asset Class", "Shares", "Price ($)", "Value ($)", "Allocation", "Target", "To Contrib"],
        table_rows,
        right_align=[2,3,4,5,6,7]
    )

    # ---------------------------------------------------------------
    # ILLUSTRATIVE MONTHLY CONTRIBUTION SCHEDULE
    # ---------------------------------------------------------------
    doc.add_heading("Illustrative Monthly Contribution Schedule", level=1)

    # Filter out holdings with zero to_contrib
    monthly_df = sec_only.copy()
    monthly_df["to_contrib_numeric"] = monthly_df["to_contrib"] # already numeric
    monthly_df = monthly_df[monthly_df["to_contrib_numeric"] > 0].copy()

    if not monthly_df.empty:
        # Use configurable monthly contribution from config
        total_monthly = TARGET_MONTHLY_CONTRIBUTION
        total_gap = monthly_df["to_contrib_numeric"].sum()
        monthly_df["monthly_contrib"] = monthly_df["to_contrib_numeric"] / total_gap * total_monthly
        monthly_df["share_of_monthly"] = monthly_df["monthly_contrib"] / total_monthly * 100

        # Build table rows
        rows = []
        for _, r in monthly_df.iterrows():
            rows.append([
                r["ticker"],
                r.get("asset_class_short", "Unknown"),
                fmt_dollar_clean(r["to_contrib_numeric"]),
                fmt_dollar_clean(r["monthly_contrib"]),
                f"{r['share_of_monthly']:.1f}%",
            ])

        add_table(
            doc,
            ["Ticker", "Asset Class", "Gap to Target", "Monthly Contrib", "Share of Monthly"],
            rows,
            right_align=[2,3,4]
        )

        # Footer paragraph
        footer = ("At approximately "
                  f"${total_monthly:,.0f}/month, this schedule allocates contributions "
                  "proportionally to each holding's gap. It would take about "
                  f"{total_gap / total_monthly:.1f} months to close all gaps, assuming flat markets.")
        doc.add_paragraph(footer)

    # ---------------------------------------------------------------
    # ASSET CLASS ALLOCATION TABLE
    # ---------------------------------------------------------------
    doc.add_heading("Asset Class Allocation", level=1)

    # Exclude CASH for asset class calculations if you want
    sec_no_cash = sec_only.copy()  # sec_only already includes CASH if needed
    total_value = sec_no_cash["value"].sum()

    # Merge short asset class and target_pct from holdings
    sec_merge = sec_no_cash.merge(
        holdings_df[["ticker", "target_pct"]],
        on="ticker",
        how="left",
        suffixes=("", "_holdings")
    )

    # Compute actual allocations per asset class
    asset_group = (
        sec_merge.groupby("asset_class_short")
        .agg(
            value=("value", "sum"),
            target_pct=("target_pct_holdings", "sum")
        )
        .reset_index()
    )

    # Compute actual percentage allocation
    asset_group["actual_pct"] = asset_group["value"] / total_value * 100

    # Compute delta
    asset_group["delta_pct"] = asset_group["actual_pct"] - asset_group["target_pct"]

    # Format columns for display
    asset_group["actual_pct_fmt"] = asset_group["actual_pct"].map(lambda x: f"{x:.2f}%")
    asset_group["target_pct_fmt"] = asset_group["target_pct"].map(lambda x: f"{x:.2f}%")
    asset_group["delta_pct_fmt"] = asset_group["delta_pct"].map(lambda x: f"+{x:.2f}%" if x >= 0 else f"{x:.2f}%")

    # Build table rows
    table_rows = []
    for _, row in asset_group.iterrows():
        table_rows.append([
            row["asset_class_short"],
            fmt_dollar_clean(row["value"]),
            row["actual_pct_fmt"],
            row["target_pct_fmt"],
            row["delta_pct_fmt"]
        ])

    # Add TOTAL row
    total_value_sum = asset_group["value"].sum()
    total_actual_pct = asset_group["actual_pct"].sum()
    total_target_pct = asset_group["target_pct"].sum()
    total_delta_pct = asset_group["delta_pct"].sum()

    table_rows.append([
        "TOTAL",
        fmt_dollar_clean(total_value_sum),
        f"{total_actual_pct:.2f}%",
        f"{total_target_pct:.2f}%",
        f"{total_delta_pct:+.2f}%"
    ])

    # Add table
    add_table(
        doc,
        ["Asset Class", "Value ($)", "Actual %", "Target %", "Delta %"],
        table_rows,
        right_align=[1, 2, 3, 4]
    )

    # Footer
    doc.add_paragraph("Delta % = actual allocation minus target allocation").paragraph_format.space_before = Pt(2)

    doc.add_page_break()


    # =============================================================
    # TICKER ALLOCATION PIE CHART
    # =============================================================
    # Build ticker allocation group
    ticker_group = sec_only.copy()
    ticker_group = ticker_group[ticker_group["value"] > 0]  # exclude zero-value tickers

    # Sort descending by value to avoid adjacent small slices
    ticker_group = ticker_group.sort_values(by="value", ascending=False)

    ticker_labels = ticker_group["ticker"].tolist()
    ticker_values = ticker_group["value"].tolist()

    img_stream = plot_ticker_allocation_pie(ticker_labels, ticker_values)

    # Insert into DOCX
    doc.add_heading("Ticker Allocation", level=1)

    # Create a paragraph for the image and center it
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(5))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Caption centered
    caption = doc.add_paragraph("Pie chart of current ticker allocations.", style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ---------------------------------------------------------------
    # TICKER ALLOCATION VS TARGET BAR CHART
    # ---------------------------------------------------------------
    
    # Use already loaded holdings_df

    # Merge actual values from sec_only with target_pct
    ticker_merge = sec_only[["ticker", "value"]].merge(
        holdings_df[["ticker", "target_pct"]],
        on="ticker",
        how="left"
    )
    ticker_merge["target_pct"] = ticker_merge["target_pct"].fillna(0.0)

    # Compute actual % allocation
    total_value = ticker_merge["value"].sum()
    ticker_merge["actual_pct"] = ticker_merge["value"] / total_value * 100

    img_stream = plot_ticker_allocation_bar(ticker_merge)

    # Insert into DOCX and center
    doc.add_heading("Ticker Allocation vs Target", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Actual vs target allocation by ticker.", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ---------------------------------------------------------------
    # ASSET CLASS ALLOCATION PIE & VS TARGET
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Asset Class Allocation Breakdown", level=1)
    
    # Use already loaded holdings_df

    # Merge actual values from sec_only
    sec_only_grouped = (
        sec_only.groupby("asset_class_short")
        .agg(value=("value", "sum"))
        .reset_index()
    )
    merged = holdings_df.groupby("asset_class")["target_pct"].sum().reset_index()
    merged["asset_class_short"] = merged["asset_class"].map({
        "US Large Cap":"US LC",
        "US Growth":"US Growth",
        "US Small Cap":"US SC",
        "International Equity":"INTL EQTY",
        "Gold / Precious Metals":"GOLD / PM",
        "Digital Assets":"DIGITAL",
        "US Bonds":"US Bonds",
        "CASH":"CASH"
    })
    ticker_merge = pd.merge(sec_only_grouped, merged[["asset_class_short","target_pct"]],
                            on="asset_class_short", how="left").fillna(0.0)

    # Sort descending by actual value
    ticker_merge = ticker_merge.sort_values("value", ascending=False).reset_index(drop=True)
    ticker_merge["actual_pct"] = ticker_merge["value"] / ticker_merge["value"].sum() * 100

    # ---------------- PIE CHART ----------------
    img_stream = plot_asset_allocation_pie(ticker_merge["asset_class_short"], ticker_merge["actual_pct"])

    # Insert into DOCX and center
    doc.add_heading("Asset Class Allocation", level=2)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(5))  # match ticker pie width
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ---------------- BAR CHART ----------------
    img_stream = plot_asset_allocation_bar(ticker_merge)

    doc.add_heading("Asset Class Allocation vs Target", level=2)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # ASSET ALLOCATION OVER TIME — STACKED AREA CHART
    # ---------------------------------------------------------------
    doc.add_heading("Asset Allocation Over Time", level=2)

    try:
        tx_hist = load_transactions_raw().copy()

        # Need both transactions and prices to build history
        if tx_hist.empty or prices.empty:
            doc.add_paragraph(
                "Insufficient transaction or price history to build an allocation-over-time chart."
            )
        else:

            # ---------------------------------------------------------------
            # Build a continuous daily index & forward-fill PV
            # ---------------------------------------------------------------
            full_index = pd.date_range(
                start=pv.index.min(),
                end=pv.index.max(),
                freq="D"
            )

            pv_mod = pv.reindex(full_index).ffill()
            alloc_index = full_index

            # Pivot transactions into daily share changes per ticker
            tx_hist["date"] = pd.to_datetime(tx_hist["date"])
            pos_changes = (
                tx_hist
                .pivot_table(
                    index="date",
                    columns="ticker",
                    values="shares",
                    aggfunc="sum"
                )
                .sort_index()
            )

            # Align positions to full daily index, FF to avoid dips
            pos_changes = pos_changes.reindex(alloc_index, fill_value=0.0)
            pos_daily = pos_changes.cumsum().ffill().bfill()

            # -------------------------------------------------------
            # Reconcile position history to current holdings
            # -------------------------------------------------------
            shares_map = (
                sec_only[["ticker", "shares"]]
                .set_index("ticker")["shares"]
                .to_dict()
            )

            # Adjust existing ticker columns to match final snapshot
            for t in list(pos_daily.columns):
                if t == "CASH":
                    continue
                if t in shares_map:
                    target = float(shares_map[t])
                    current = float(pos_daily[t].iloc[-1])
                    diff = target - current
                    if abs(diff) > 1e-8:
                        pos_daily[t] = pos_daily[t] + diff

            # Add static columns for tickers never seen in tx_hist
            for t, shares in shares_map.items():
                if t == "CASH":
                    continue
                if t not in pos_daily.columns and t in prices.columns:
                    pos_daily[t] = float(shares)

            # Restrict to tickers we actually have prices for
            common_tickers = [t for t in pos_daily.columns if t in prices.columns]

            if not common_tickers:
                doc.add_paragraph(
                    "No overlapping tickers between transactions and price history."
                )
            else:
                pos_daily = pos_daily[common_tickers]

                # Forward-fill prices on full daily index
                px_aligned = (
                    prices[common_tickers]
                    .reindex(alloc_index)
                    .ffill()
                    .bfill()
                )

                # Daily market value per ticker
                mv_daily = pos_daily * px_aligned

                # Map tickers → asset classes
                holdings_map = (
                    holdings_df[["ticker", "asset_class"]]
                    .drop_duplicates()
                    .set_index("ticker")["asset_class"]
                    .to_dict()
                )

                def map_asset_class_short(ticker: str) -> str:
                    full = holdings_map.get(ticker, "Unknown")
                    return asset_class_map.get(full, full)

                mv_daily.columns = [map_asset_class_short(t) for t in mv_daily.columns]

                # Aggregate by asset class
                mv_by_class = mv_daily.T.groupby(level=0).sum().T

                # Align PV
                pv_mod_aligned = pv_mod.reindex(mv_by_class.index).ffill().bfill()
                invested_total = mv_by_class.sum(axis=1)
                cash_series = pv_mod_aligned - invested_total
                mv_by_class["Cash"] = cash_series

                # Convert to allocation percentages
                total_mv = mv_by_class.sum(axis=1).replace(0, np.nan)
                alloc = mv_by_class.div(total_mv, axis=0).dropna(how="all")

                if alloc.empty:
                    doc.add_paragraph(
                        "No valid allocation history could be computed."
                    )
                else:
                    alloc_pct = alloc * 100.0

                    img_stream = plot_asset_allocation_history(alloc_pct)

                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(7))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cap = doc.add_paragraph(
                        "Figure: Daily asset class allocation over time, including cash.",
                        style="Normal",
                    )
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception:
        doc.add_paragraph(
            "Asset allocation over time chart could not be generated.",
            style="Normal"
        )


    # ---------------------------------------------------------------
    # PORTFOLIO VS BENCHMARKS — MTD & YTD (PRICE RETURN)
    # ---------------------------------------------------------------

    # Keep ORIGINAL pv for horizon logic (TWR + MD + snapshot consistency)
    pv_trading = pv.sort_index()

    # Build a DAILY PV series for CHARTS ONLY
    pv_daily = pv_trading.reindex(
        pd.date_range(pv_trading.index.min(), pv_trading.index.max(), freq="D")
    ).ffill()

    portfolio_value = pv_daily.copy()



    # Current "as of" date
    as_of = portfolio_value.index[-1]

    # ---------------- MTD (Institutional: last trading day of prior month) ----------------
    pv_dates = pv.index
    mtd_start = get_portfolio_horizon_start(pv, inception_date, "MTD")
    if mtd_start is None:
        mtd_start = pv.index.min()

    # unified anchor for EVERYTHING (portfolio + benchmarks)
    pv_nonzero = pv[pv > 0]
    twr_anchor = max(mtd_start, pv_nonzero.index.min())

    portfolio_mtd = portfolio_value[portfolio_value.index >= twr_anchor]

    portfolio_mtd = (portfolio_mtd / portfolio_mtd.iloc[0] - 1) * 100


    # ---------------- YTD ----------------
    ytd_start = get_portfolio_horizon_start(pv, inception_date, "YTD")
    if ytd_start is None:
        ytd_start = pv.index.min()

    portfolio_ytd = portfolio_value[portfolio_value.index >= ytd_start]

    portfolio_ytd = (portfolio_ytd / portfolio_ytd.iloc[0] - 1) * 100

    # Define benchmark tickers and fetch historical prices
    benchmarks = {
        "S&P 500 (SPY)": "SPY",
        "AOR (Global 60/40)": "AOR",
        "AOK (Conservative 40/60)": "AOK"
    }

    benchmark_prices = {}
    for name, ticker in benchmarks.items():
        hist = fetch_price_history([ticker])
        hist.index = pd.to_datetime(hist.index)
        benchmark_prices[name] = hist[ticker]

    # Slice and compute MTD/YTD for benchmarks
    benchmark_returns = {}
    for name, series in benchmark_prices.items():
        # MTD (MUST USE twr_anchor)
        mtd_series = series[series.index >= twr_anchor]
        mtd_cum = (mtd_series / mtd_series.iloc[0] - 1) * 100

        # YTD
        ytd_series = series[series.index >= ytd_start]
        ytd_cum = (ytd_series / ytd_series.iloc[0] - 1) * 100
        benchmark_returns[name] = {"MTD": mtd_cum, "YTD": ytd_cum}

    # ---------------- MTD TWR-BASED CHART ----------------
    doc.add_page_break()
    doc.add_heading("Portfolio vs Benchmarks Analysis", level=1)
    doc.add_heading("MTD Cumulative Return — Portfolio vs Benchmarks", level=2)

    # 1) Portfolio TWR cumulative curve (based on pv)
    mtd_start_twr = twr_anchor

    # Correct GIPS-compliant daily TWR (flows at start of day)
    cf_ext_local = load_cashflows_external().copy()
    daily_ret = []

    pv_dates = pv.index
    for i in range(1, len(pv_dates)):
        d0 = pv_dates[i-1]
        d1 = pv_dates[i]

        flow = cf_ext_local.loc[cf_ext_local["date"] == d1, "amount"].sum()
        denom = pv.loc[d0] + flow
        if denom <= 0:
            continue

        R = (pv.loc[d1] - denom) / denom
        daily_ret.append((d1, R))


    # Convert to cumulative TWR curve
    twr_curve = pd.Series(1.0, index=[d for d,_ in daily_ret])
    running = 1.0
    for d,R in daily_ret:
        running *= (1 + R)
        twr_curve.loc[d] = running

    # Unified since-inception anchor for both SI chart and SI benchmark table
    si_start = inception_date


    # Slice only the MTD part
    twr_mtd = twr_curve[twr_curve.index >= mtd_start_twr]
    twr_mtd = (twr_mtd / twr_mtd.iloc[0] - 1) * 100


    # 2) Benchmarks: ALSO anchored to the same start date as portfolio
    benchmark_map = {
        "S&P 500 (SPY)": "SPY",
        "AOR (Global 60/40)": "AOR",
        "AOK (Conservative 40/60)": "AOK"
    }

    benchmark_curves_mtd = {}

    for name, ticker in benchmark_map.items():
        try:
            hist = fetch_price_history([ticker])
            hist.index = pd.to_datetime(hist.index)
            ser = hist[ticker]

            # Align to portfolio’s actual curve start
            ser = ser[ser.index >= mtd_start_twr]

            if len(ser) > 1:
                ser_norm = (ser / ser.iloc[0] - 1.0) * 100.0
                benchmark_curves_mtd[name] = ser_norm
        except:
            pass

    # ---------------- UPDATED MTD CHART (Option C sizing + matching fonts) ----------------
    img_stream = plot_mtd_cumulative_return(twr_mtd, benchmark_curves_mtd)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6.5))   # <<< bigger but not full-page
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_paragraph(
        "Figure: Portfolio returns use true TWR based on a flow-adjusted, end-of-day PV series (external flows treated at start-of-day). Benchmark returns use simple price return and are rebased to the same start date for comparability.",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    # ---------------- SINCE INCEPTION CHART (TWR vs Benchmarks) ----------------

    # ---- build SI pct curve BEFORE any x-axis work ----
    twr_si_curve = twr_curve.copy()
    twr_si_pct = (twr_si_curve - 1.0) * 100.0
    si_idx = twr_si_pct.index

    # Benchmarks since inception — price return rebased to si_start
    benchmark_curves_si = {}

    for name, ticker in benchmark_map.items():
        try:
            hist = fetch_price_history([ticker])
            hist.index = pd.to_datetime(hist.index)
            ser = hist[ticker].dropna()

            # Align to portfolio SI start
            ser = ser[ser.index >= si_start]

            if len(ser) > 1:
                ser_norm = (ser / ser.iloc[0] - 1.0) * 100.0
                benchmark_curves_si[name] = ser_norm
        except Exception:
            continue

    img_stream = plot_si_cumulative_return(twr_si_pct, benchmark_curves_si, si_idx)

    doc.add_heading("Since Inception Cumulative Return — Portfolio vs Benchmarks", level=2)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6.5))   # MATCH MTD WIDTH EXACTLY
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Figure: Portfolio curve uses true since-inception TWR based on the flow-adjusted PV series. "
        "Benchmarks use simple price returns rebased to the same since-inception start date.",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # ---------------------------------------------------------------
    # PERFORMANCE VS BENCHMARKS (ALL HORIZONS — TWR CONSISTENT)
    # ---------------------------------------------------------------
    doc.add_heading("Performance vs Benchmarks (All Horizons)", level=2)

    # Horizons to show (must match your snapshot / twr_df + SI)
    horizons_all = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "SI"]

    # Portfolio TWR (numeric, decimals) from twr_raw
    port_ret_num = {row["Horizon"]: row["Return"] for _, row in twr_raw.iterrows()}
    port_ret_str = {row["Horizon"]: row["Return"] for _, row in twr_df.iterrows()}

    # Inject SI (since inception) into both maps
    if not pd.isna(twr_si):
        port_ret_num["SI"] = twr_si
        # use non-annualized SI TWR here for the table
        port_ret_str["SI"] = fmt_pct_clean(twr_si)

    # Benchmarks and their tickers (keys used for lookups below)
    bm_defs = {
        "S&P 500 %": "SPY",
        "Global 60/40 %": "AOR",
        "Conservative 40/60 %": "AOK",
    }

    # Cache benchmark price history once
    bm_prices = {}
    for col_label, ticker in bm_defs.items():
        try:
            hist = fetch_price_history([ticker])
            ser = hist[ticker].dropna()
        except Exception:
            ser = pd.Series(dtype=float)
        bm_prices[col_label] = ser

    # Benchmark return as a decimal, rebased to same start as portfolio
    def compute_bm_ret_decimal(col_label: str, horizon: str) -> float:
        ser = bm_prices[col_label]
        if ser.empty:
            return np.nan

        if horizon == "SI":
            start = si_start  # same since-inception start used in the SI chart
        else:
            start = get_portfolio_horizon_start(pv, inception_date, horizon)

        if start is None:
            return np.nan

        ser_h = ser[ser.index >= start]
        if len(ser_h) < 2:
            return np.nan

        return ser_h.iloc[-1] / ser_h.iloc[0] - 1.0  # decimal

    # Build table rows with excess return columns
    combined_rows = []
    for h in horizons_all:
        port_str = port_ret_str.get(h, "N/A")
        port_dec = port_ret_num.get(h, np.nan)

        sp_dec   = compute_bm_ret_decimal("S&P 500 %", h)
        g6040_dec = compute_bm_ret_decimal("Global 60/40 %", h)
        c4060_dec = compute_bm_ret_decimal("Conservative 40/60 %", h)

        # Benchmark display strings
        sp_str    = fmt_pct_clean(sp_dec)
        g6040_str = fmt_pct_clean(g6040_dec)
        c4060_str = fmt_pct_clean(c4060_dec)

        # Excess = portfolio − benchmark (all decimals)
        if pd.isna(port_dec) or pd.isna(sp_dec):
            sp_excess_str = "N/A"
        else:
            sp_excess_str = fmt_pct_clean(port_dec - sp_dec)

        if pd.isna(port_dec) or pd.isna(g6040_dec):
            g6040_excess_str = "N/A"
        else:
            g6040_excess_str = fmt_pct_clean(port_dec - g6040_dec)

        if pd.isna(port_dec) or pd.isna(c4060_dec):
            c4060_excess_str = "N/A"
        else:
            c4060_excess_str = fmt_pct_clean(port_dec - c4060_dec)

        combined_rows.append([
            h,
            port_str,
            sp_str,
            sp_excess_str,
            g6040_str,
            g6040_excess_str,
            c4060_str,
            c4060_excess_str,
        ])

    # Shorten last header label to stop wrapping
    headers = [
        "Horizon",
        "Portfolio %",
        "S&P 500 %",
        "Excess",
        "Global 60/40 %",
        "Excess",
        "Cons 40/60 %",
        "Excess",
    ]

    add_table(
        doc,
        headers,
        combined_rows,
        right_align=[1, 2, 3, 4, 5, 6, 7]
    )
    
    # ---- FIX WIDTHS ONLY FOR THIS BENCHMARKS TABLE ----
    table = doc.tables[-1]          # get last added table
    table.autofit = False
    table.allow_autofit = False

    col_widths = [
        Inches(0.9),  # Horizon
        Inches(0.9),  # Portfolio %
        Inches(0.9),  # S&P 500 %
        Inches(0.9),  # Excess vs S&P
        Inches(1.15),  # Global 60/40 %
        Inches(0.9),  # Excess vs 60/40
        Inches(1.1),  # Cons 40/60 %
        Inches(0.9),  # Excess vs 40/60
    ]

    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w


    doc.add_paragraph(
        "Portfolio % uses flow-adjusted TWR. Benchmark % uses simple price returns "
        "rebased to the same start date for each horizon. Excess = Portfolio % minus benchmark % for the same horizon.",
        style="Normal"
    )

    # -----------------------------------------------------------
    # EXCESS RETURN VS BENCHMARKS (BY HORIZON) — FINAL VERSION
    # -----------------------------------------------------------
    doc.add_heading("Excess Return vs Benchmarks (by Horizon)", level=3)

    horizons_plot = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "SI"]
    bm_labels = list(bm_defs.keys())  # e.g. ["S&P 500 %", "Global 60/40 %", "Conservative 40/60 %"]

    # --- Local helper ---
    def compute_bm_ret(col_label: str, horizon: str) -> float:
        ser = bm_prices.get(col_label)
        if ser is None or ser.empty:
            return np.nan
        if horizon == "SI":
            start = si_start
        else:
            start = get_portfolio_horizon_start(pv, inception_date, horizon)
        if start is None:
            return np.nan
        ser_h = ser[ser.index >= start]
        if len(ser_h) < 2:
            return np.nan
        return ser_h.iloc[-1] / ser_h.iloc[0] - 1.0  # decimal return

    # --- Build numeric excess matrix (Portfolio − Benchmark) ---
    excess = {bm: [] for bm in bm_labels}

    for h in horizons_plot:
        p_val = port_ret_num.get(h, np.nan)
        for bm in bm_labels:
            b_val = compute_bm_ret(bm, h)
            if pd.isna(p_val) or pd.isna(b_val):
                excess[bm].append(np.nan)
            else:
                excess[bm].append((p_val - b_val) * 100.0)  # percent points

    img_stream = plot_excess_return(horizons_plot, bm_labels, excess)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6.4))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph(
        "Positive bars show outperformance; negative bars show underperformance.",
        style="Normal",
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


    # =============================================================
    # NEW: SINCE-INCEPTION (SI) RETURNS FOR TICKERS & ASSET CLASSES
    # =============================================================

    # We want a true SI return per ticker (since first trade),
    # and a value-weighted rollup for each asset class.
    si_return_map = {}

    tx_all = load_transactions_raw().copy()
    as_of_port = pv.index.max()

    # Use the same price history we already fetched for tickers
    # (non-CASH only; CASH gets 0%).
    for t in sec_only["ticker"].unique():
        if t == "CASH":
            si_return_map[t] = 0.0
            continue

        if t not in prices.columns:
            # No price history → treat as 0% to avoid N/A spam
            si_return_map[t] = 0.0
            continue

        price_series = prices[t].dropna()
        if price_series.empty:
            si_return_map[t] = 0.0
            continue

        tx_t = tx_all[tx_all["ticker"] == t].copy()
        if tx_t.empty:
            # No transactions recorded → fallback to 0%
            si_return_map[t] = 0.0
            continue

        tx_t = tx_t.sort_values("date")
        first_trade = tx_t["date"].min()

        # End date = min(portfolio PV max date, price series max date)
        as_of_price = price_series.index.max()
        end = min(as_of_port, as_of_price)

        if end <= first_trade:
            si_return_map[t] = 0.0
            continue

        # Run MD from first trade date
        try:
            si_ret = modified_dietz_for_ticker_window(
                t,
                price_series,
                tx_t,
                first_trade,
                end,
            )
        except Exception:
            si_ret = np.nan

        # If MD blows up to NaN, leave it as NaN so it shows as "N/A"
        if pd.isna(si_ret):
            si_return_map[t] = np.nan
        else:
            si_return_map[t] = float(si_ret)


    # Attach SI returns to security-level table as a DECIMAL,
    # then we will format for display in the horizon table only.
    if not sec_full.empty:
        sec_full["SI"] = sec_full["ticker"].map(lambda t: si_return_map.get(t, 0.0))

    # Now roll these up to asset classes, value-weighted by current MV.
    if not class_full.empty:
        # Merge current market value into a helper frame
        sec_mv = sec_only[["ticker", "asset_class", "value"]].copy()
        sec_mv["SI"] = sec_mv["ticker"].map(lambda t: si_return_map.get(t, 0.0))

        si_by_class = {}
        for ac, grp in sec_mv.groupby("asset_class"):
            grp = grp.dropna(subset=["value"])
            if grp.empty:
                si_by_class[ac] = 0.0
                continue

            # Only tickers with non-null SI
            sub = grp.dropna(subset=["SI"])
            if sub.empty:
                si_by_class[ac] = 0.0
                continue

            w = sub["value"] / sub["value"].sum()
            si_by_class[ac] = float((w * sub["SI"]).sum())

        # Add SI as decimal to class_full
        class_full["SI"] = class_full["asset_class"].map(
            lambda ac: si_by_class.get(ac, 0.0)
        )

    # Format ticker-level SI for display in the horizon table
    if not sec_full.empty and "SI" in sec_full.columns:
        sec_full["SI"] = sec_full["SI"].apply(fmt_pct_clean)


    # =============================================================
    # MULTI-HORIZON RETURNS BY ASSET CLASS & TICKER (MODIFIED DIETZ)
    # =============================================================
    doc.add_heading("Horizon Performance Tables", level=1)

    doc.add_heading("Returns Grouped by Asset Class (Modified Dietz)", level=2)

    horizon_cols = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "SI"]

    if class_full.empty and sec_full.empty:
        doc.add_paragraph("No security-level cashflow data available.")
    else:
        # Make a formatted view of class_full so we don't double-multiply anything
        class_view = class_full.copy()
        for col in horizon_cols:
            if col in class_view.columns and col != "SI":
                # Non-SI horizons: standard formatting
                class_view[col] = class_view[col].apply(fmt_pct_clean)

        # SI is a decimal; format it once here so asset-class SI shows as %.
        if "SI" in class_view.columns:
            class_view["SI"] = class_view["SI"].apply(fmt_pct_clean)


        # Order asset classes based on class_full
        asset_classes = (
            class_view["asset_class"]
            .dropna()
            .drop_duplicates()
            .tolist()
        )

        rows = []

        for ac in asset_classes:
            # ----- Asset class row -----
            class_row = class_view[class_view["asset_class"] == ac]
            if not class_row.empty:
                r = class_row.iloc[0]
                row_vals = [asset_class_map.get(ac, ac)]
                for h in horizon_cols:
                    row_vals.append(safe(r.get(h)))
                rows.append(row_vals)

            # ----- Ticker rows under that class -----
            tickers_in_ac = (
                sec_full[sec_full["asset_class"] == ac]
                .sort_values("ticker")["ticker"]
                .tolist()
            )

            for t in tickers_in_ac:
                tr = sec_full[sec_full["ticker"] == t].iloc[0]
                t_vals = [f"  {t}"]
                for h in horizon_cols:
                    t_vals.append(safe(tr.get(h)))
                rows.append(t_vals)

        add_table(
            doc,
            ["Asset Class / Ticker"] + horizon_cols,
            rows,
            right_align=list(range(1, 1 + len(horizon_cols)))
        )
        
        # ---- FIX WIDTHS ONLY FOR THIS HORIZON RETURNS TABLE ----
        table = doc.tables[-1]  # get last added table
        table.autofit = False
        table.allow_autofit = False

        col_widths = [
            Inches(1.4),  # Asset Class / Ticker
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8),
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)
        ]

        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w


        doc.add_paragraph(
            "Returns above are computed using a per-holding and per-asset-class Modified Dietz methodology for each horizon.",
            style="Normal"
        )

    # =============================================================
    # MULTI-HORIZON P/L BY ASSET CLASS & TICKER (ECONOMIC, MD-CONSISTENT)
    # =============================================================

    doc.add_heading("P/L Grouped by Asset Class", level=2)

    horizons_pl = ["1D", "1W", "MTD", "1M", "3M", "6M", "1Y", "SI"]

    if class_full.empty and sec_full.empty:
        doc.add_paragraph("No security-level cashflow data available.")
    else:

        def compute_asset_class_pl(asset_class: str, h: str):
            """
            Economic P/L for an asset class over horizon h.

            Computed as the sum of ticker-level economic P/L for all tickers
            in that asset class, where each ticker P/L is already horizon-
            aligned to the same get_horizon_start() logic.

            This is additive and therefore fully consistent with the ticker-
            level Modified Dietz construction.
            """
            # Use sec_only for the universe of tickers and asset_class mapping
            tickers_in_ac = (
                sec_only[sec_only["asset_class"] == asset_class]["ticker"]
                .dropna()
                .unique()
            )

            total_pl = 0.0
            seen_any = False

            for t in tickers_in_ac:
                pl_str = get_ticker_pl_str(t, h)
                if pl_str in (None, "N/A"):
                    continue
                try:
                    # parse formatted "$x,xxx.xx" back to float
                    val = float(pl_str.replace("$", "").replace(",", ""))
                except Exception:
                    continue
                total_pl += val
                seen_any = True

            if not seen_any:
                return "N/A"
            return fmt_dollar_clean(total_pl)

        rows_pl = []

        # Use the same asset class ordering as the returns table
        asset_classes_pl = (
            class_full["asset_class"]
            .dropna()
            .drop_duplicates()
            .tolist()
        )

        for ac in asset_classes_pl:
            # ----- Asset class row -----
            row_vals = [asset_class_map.get(ac, ac)]
            for h in horizons_pl:
                row_vals.append(compute_asset_class_pl(ac, h))
            rows_pl.append(row_vals)

            # ----- Ticker rows under that class -----
            tickers_in_ac = (
                sec_only[sec_only["asset_class"] == ac]
                .sort_values("ticker")["ticker"]
                .tolist()
            )

            for t in tickers_in_ac:
                t_vals = [f"  {t}"]
                for h in horizons_pl:
                    t_vals.append(get_ticker_pl_str(t, h))
                rows_pl.append(t_vals)

        # -------------------------------------------------------
        # FIX: Add Cash / Reconciliation Row
        # -------------------------------------------------------
        cash_row = ["Cash / Reconciliation"]
        
        for h in horizons_pl:
            # 1. Get Portfolio Total P/L (External Flows)
            if h == "SI":
                port_pl = pl_si # calculated in run_engine
            else:
                port_pl = calculate_horizon_pl(pv, inception_date, cf_ext, h)
                
            if port_pl is None:
                cash_row.append("N/A")
                continue
                
            # 2. Sum Ticker P/Ls (Internal Flows)
            sum_ticker_pl = 0.0
            
            # Iterate all non-cash tickers
            all_tickers = sec_only[sec_only["ticker"] != "CASH"]["ticker"].unique()
            for t in all_tickers:
                as_of_dt = pv.index.max()
                if h == "SI":
                    raw_start = None
                else:
                    raw_start = get_portfolio_horizon_start(pv, inception_date, h)
                    
                val = calculate_ticker_pl(t, h, prices, as_of_dt, tx_all_raw, sec_only, raw_start)
                if val is not None:
                    sum_ticker_pl += val
                    
            # 3. Diff is Cash P/L (captures interest, FX, drag, etc.)
            diff = port_pl - sum_ticker_pl
            cash_row.append(fmt_dollar_clean(diff))
            
        rows_pl.append(cash_row)

        add_table(
            doc,
            ["Asset Class / Ticker"] + horizons_pl,
            rows_pl,
            right_align=list(range(1, 1 + len(horizons_pl)))
        )
        
        # ---- FIX WIDTHS ONLY FOR THIS HORIZON P/L TABLE ----
        table = doc.tables[-1]  # get last added table
        table.autofit = False
        table.allow_autofit = False

        col_widths = [
            Inches(1.4),  # Asset Class / Ticker
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8),
            Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)
        ]

        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w


        doc.add_paragraph(
            "P/L above is economic P/L by horizon (MV_end − MV_start − net internal flows), "
            "rolled up additively from tickers to asset classes using the same horizon definitions as the Modified Dietz returns.",
            style="Normal"
        )


    # ---------------------------------------------------------------
    # DETAILED CASH FLOW REPORT (Since Inception)
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Detailed Cash Flow Report", level=1)

    # ===============================================================
    # EXTERNAL CASH FLOWS (Deposits & Withdrawals)
    # ===============================================================
    doc.add_heading("External Cash Flows (Since Inception)", level=2)

    cf_all = load_cashflows_external().copy()
    cf_all = cf_all.sort_values("date")

    if cf_all.empty:
        doc.add_paragraph("No external cash flows recorded.")
    else:
        # Total deposits, withdrawals, net
        total_deposits = cf_all.loc[cf_all["amount"] > 0, "amount"].sum()
        total_withdrawals = cf_all.loc[cf_all["amount"] < 0, "amount"].sum()
        net_ext = cf_all["amount"].sum()

        most_recent = cf_all["date"].max().strftime("%Y-%m-%d") if not cf_all.empty else "N/A"

        ext_rows = [
            ["Total Deposits", fmt_dollar_clean(total_deposits)],
            ["Total Withdrawals", fmt_dollar_clean(total_withdrawals)],
            ["Net External Flow", fmt_dollar_clean(net_ext)],
            ["Most Recent External Flow", most_recent],
        ]

        add_table(
            doc,
            ["Metric", "Value"],
            ext_rows,
            right_align=[1]
        )

    # Small space before internal section
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # ===============================================================
    # INTERNAL TRADING SUMMARY (Buys, Sells, Net, Income)
    # ===============================================================
    doc.add_heading("Internal Trading Summary (Since Inception)", level=2)

    tx_all = load_transactions_raw().copy()
    div_all = load_dividends().copy()

    if tx_all.empty and div_all.empty:
        doc.add_paragraph("No internal trading or income activity available.")
    else:
        # 1) Aggregate buys/sells
        if tx_all.empty:
            summary = pd.DataFrame(columns=["ticker", "buys", "sells"])
        else:
            summary = (
                tx_all.groupby("ticker")["amount"]
                .agg([
                    ("buys", lambda s: s[s < 0].sum()),
                    ("sells", lambda s: s[s > 0].sum())
                ])
                .reset_index()
            )

        # 2) Aggregate dividends (income)
        if div_all.empty:
            div_agg = pd.DataFrame(columns=["ticker", "income"])
        else:
            div_agg = (
                div_all.groupby("ticker")["amount"]
                .sum()
                .reset_index()
                .rename(columns={"amount": "income"})
            )

        # 3) Merge
        if summary.empty and not div_agg.empty:
            final = div_agg
            final["buys"] = 0.0
            final["sells"] = 0.0
        elif not summary.empty:
            final = summary.merge(div_agg, on="ticker", how="outer")
            final = final.fillna(0.0)
        else:
            final = pd.DataFrame(columns=["ticker", "buys", "sells", "income"])

        # 4) Compute Net
        if "income" not in final.columns:
            final["income"] = 0.0
        if "buys" not in final.columns:
            final["buys"] = 0.0
        if "sells" not in final.columns:
            final["sells"] = 0.0

        final["net"] = final["buys"] + final["sells"] + final["income"]
        final = final.sort_values("ticker")

        # Format rows for display
        rows = []
        for _, r in final.iterrows():
            rows.append([
                r["ticker"],
                fmt_dollar_clean(r["buys"]),
                fmt_dollar_clean(r["sells"]),
                fmt_dollar_clean(r["income"]),
                fmt_dollar_clean(r["net"])
            ])

        add_table(
            doc,
            ["Ticker", "Total Buys (Cash Out)", "Total Sells (Cash In)", "Total Income", "Net (Cash Flow)"],
            rows,
            right_align=[1, 2, 3, 4]
        )

    # ---------------------------------------------------------------
    # INTERNAL TRADING FLOWS — STACKED BY ASSET CLASS (SINCE INCEPTION)
    # ---------------------------------------------------------------
    doc.add_heading("Internal Trading Flows by Asset Class (Since Inception)", level=2)

    tx_raw = load_transactions_raw().copy()

    if tx_raw.empty:
        doc.add_paragraph("No internal buy/sell activity recorded.")
    else:
        # Map tickers → asset class
        asset_map = (
            holdings_df[["ticker", "asset_class"]]
            .set_index("ticker")["asset_class"]
            .to_dict()
        )
        tx_raw["asset_class"] = tx_raw["ticker"].map(asset_map).fillna("Exited Holdings")

        # Net flows per class
        net_by_class = (
            tx_raw.groupby("asset_class")["amount"]
            .sum()
            .sort_values(ascending=True)
        )

        img_stream = plot_internal_trading_flows(net_by_class)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_stream, width=Inches(6))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # PV MOUNTAIN CHART (Normalized to % Return Since Inception)
    # ---------------------------------------------------------------
    doc.add_heading("Portfolio Value Analysis", level=1)

    doc.add_heading("Portfolio Cumulative Return (Time-Weighted)", level=2)

    # FIX: Use calculated TWR series (twr_si_pct) instead of raw PV
    # This ensures deposits/withdrawals do not distort the performance chart.
    
    img_stream = plot_pv_mountain(twr_si_pct)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Figure: Cumulative Time-Weighted Return (Growth of $100). "
        "This metric isolates investment performance from external cash flows.",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # DAILY ΔPV ATTRIBUTION — EXTERNAL FLOWS vs MARKET MOVE
    # ---------------------------------------------------------------
    doc.add_heading("Daily ΔPV Attribution", level=2)

    # Build daily PV (fill missing days)
    pv_trading = pv.sort_index().copy()
    pv_daily = pv_trading.reindex(
        pd.date_range(pv_trading.index.min(), pv_trading.index.max(), freq="D")
    ).ffill()

    # Daily ΔPV
    dpv = pv_daily.diff().fillna(0.0)

    # ---------------- EXTERNAL FLOWS ----------------
    cf_all = load_cashflows_external().copy()
    if not cf_all.empty:
        cf_all = cf_all.sort_values("date")
        ext_series = (
            cf_all.groupby("date")["amount"]
            .sum()
            .reindex(pv_daily.index, fill_value=0.0)
        )
    else:
        ext_series = pd.Series(0.0, index=pv_daily.index)

    # ---------------- MARKET EFFECT ----------------
    # Residual after removing EXTERNAL flows only.
    # Internal trades (buys/sells) are pure reallocations between
    # cash and positions and should not change total PV.
    mkt_series = dpv - ext_series


    # Last 30 days window
    window_days = 30
    end_date = pv_daily.index.max()
    start_date = max(end_date - pd.Timedelta(days=window_days - 1), pv_daily.index.min())

    mask = (pv_daily.index >= start_date) & (pv_daily.index <= end_date)
    dates_win = pv_daily.index[mask]
    dpv_win = dpv[mask]
    ext_win = ext_series[mask]
    mkt_win = mkt_series[mask]

    # Cumulative ΔPV (rebased to 0)
    # ΔPV is exactly External + Market in this construction.
    cum_total = (ext_series + mkt_series).cumsum()

    cum_win = cum_total[mask] - cum_total[mask].iloc[0]

    img_stream = plot_daily_dpv_attribution(dates_win, ext_win, mkt_win, cum_win, dpv_win)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph(
        "Figure: Daily decomposition of portfolio value changes into external flows "
        "and market movements. Stacked bars sum to total ΔPV; the dashed line shows "
        "cumulative ΔPV over the same window. Y-axis is auto-centered on zero.",
        style="Normal",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---------------------------------------------------------------
    # LONG-TERM PROJECTION SCENARIOS (20 YEARS)
    # ---------------------------------------------------------------
    doc.add_heading("20-Year Projection Scenarios", level=1)

    # Load contribution amount from config
    monthly_contrib = TARGET_MONTHLY_CONTRIBUTION

    # Initial portfolio value = current PV
    initial_value = float(pv.iloc[-1])

    # Projection assumptions
    years = [1, 5, 10, 15, 20]
    rates = [0.05, 0.07, 0.09]

    # Build projection table rows
    proj_rows = []
    for yr in years:
        row = [yr]
        # Lump-only
        for r in rates:
            row.append(f"${fv_lump(initial_value, r, yr):,.0f}")
        # Lump + monthly contributions
        for r in rates:
            fv_total = fv_lump(initial_value, r, yr) + fv_contrib(monthly_contrib, r, yr)
            row.append(f"${fv_total:,.0f}")
        proj_rows.append(row)

    # Build table
    headers = [
        "Year",
        "5% (Lump)",
        "7% (Lump)",
        "9% (Lump)",
        f"5% (+${monthly_contrib:,.0f}/mo)",
        f"7% (+${monthly_contrib:,.0f}/mo)",
        f"9% (+${monthly_contrib:,.0f}/mo)"
    ]

    add_table(
        doc,
        headers,
        proj_rows,
        right_align=[1,2,3,4,5,6]
    )

    # Add spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)


    # ---------------------------------------------------------------
    # LONG-TERM PROJECTION CHART
    # ---------------------------------------------------------------

    # Year-by-year curves for full plotting
    horizon_years = range(0, 21)

    proj_lump = {r: [] for r in rates}
    proj_contrib = {r: [] for r in rates}

    for yr in horizon_years:
        for r in rates:
            proj_lump[r].append(fv_lump(initial_value, r, yr))
            proj_contrib[r].append(fv_lump(initial_value, r, yr) + fv_contrib(monthly_contrib, r, yr))

    img_stream = plot_long_term_projection(horizon_years, proj_lump, proj_contrib, rates, monthly_contrib)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph("Figure: Long-term projections.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------------
    # RISK & VOLATILITY ANALYSIS (APPENDIX)
    # ---------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Risk & Volatility Analysis", level=1)

    # ===========================
    # Expected Volatility By Asset Class (Bar Chart)
    # ===========================
    doc.add_heading("Expected Volatility by Asset Class", level=2)

    vol_data = {
        "Digital Assets": 75,
        "Gold / Precious Metals": 14,
        "International Equity": 17.5,
        "US Bonds": 6,
        "US Growth": 19,
        "US Large Cap": 15,
        "US Small Cap": 24
    }

    classes = list(vol_data.keys())
    vol_vals = list(vol_data.values())

    img_stream = plot_expected_volatility(classes, vol_vals)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph("Figure: Approximate volatility estimate.", style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===========================
    # Risk vs Expected Return (Scatter)
    # ===========================
    doc.add_heading("Risk vs Expected Return", level=2)

    exp_return_data = {
        "Digital Assets": (75, 12),
        "US Small Cap": (24, 9.0),
        "US Growth": (19, 8.5),
        "International Equity": (17.5, 8.0),
        "US Large Cap": (15, 7.0),
        "Gold / Precious Metals": (14.5, 5.5),
        "US Bonds": (6, 4.5)
    }

    vols = [v[0] for v in exp_return_data.values()]
    rets = [v[1] for v in exp_return_data.values()]
    labels = list(exp_return_data.keys())

    img_stream = plot_risk_return(vols, rets, labels)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph("Figure: Trade-off between expected return and volatility.", style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ---------------------------------------------------------------
    # SECTOR ALLOCATION HEATMAP
    # ---------------------------------------------------------------
    doc.add_page_break()

    # Hardcode ETF sector map from config
    sector_exposure = defaultdict(float)

    # Sector normalization to avoid duplicates
    SECTOR_NORMALIZATION = {
        "Comm Services": "Communication Services",
        "Consumer Disc.": "Consumer Discretionary",
        "Information Technology": "Tech",
        "Other": None,  # ignore "Other" buckets
    }

    # Use raw market values and exclude CASH from the denominator so that
    # sector weights are based on invested assets only.
    sector_universe = sec_only[sec_only["ticker"].isin(ETF_SECTOR_MAP.keys())].copy()
    sector_universe = sector_universe[sector_universe["value"] > 0]

    total_invested = sector_universe["value"].sum()

    if total_invested > 0:
        for _, row in sector_universe.iterrows():
            ticker = row["ticker"]
            # weight of this holding as % of invested (non-CASH) assets
            weight_pct = (row["value"] / total_invested) * 100.0

            etf_sectors = ETF_SECTOR_MAP.get(ticker, {})
            for sector, pct in etf_sectors.items():
                norm_sector = SECTOR_NORMALIZATION.get(sector, sector)
                if norm_sector is None:
                    continue
                sector_exposure[norm_sector] += weight_pct * pct / 100.0


    # Convert to sorted DataFrame
    sector_df = pd.DataFrame(
        list(sector_exposure.items()),
        columns=["Sector", "Exposure"]
    ).sort_values("Exposure", ascending=True)

    img_stream = plot_sector_allocation(sector_df)

    doc.add_heading("Sector Allocation", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure: Approximate sector exposure for portfolio ETFs.", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =============================================================
    # METHODOLOGY APPENDIX (CLEAN ONE-PAGER)
    # =============================================================
    doc.add_page_break()
    doc.add_heading("Appendix — Methodology Summary", level=1)

    # Subtle intro line
    p = doc.add_paragraph()
    r = p.add_run("This page summarizes how the figures in this report are calculated.")
    r.italic = True

    # --- Data & Portfolio Value ---
    p = doc.add_paragraph()
    r = p.add_run("Data and portfolio value. ")
    r.bold = True
    p.add_run(
        "Holdings, transactions, and cash flows are taken from the source files and combined with daily market prices "
        "to build a consistent portfolio value series. External cash flows (deposits and withdrawals) are separated "
        "from internal trading activity (buys and sells) so that performance can be evaluated independently of funding decisions."
    )

    # --- Portfolio Time-Weighted Returns (TWR) ---
    p = doc.add_paragraph()
    r = p.add_run("Portfolio time-weighted returns (TWR). ")
    r.bold = True
    p.add_run(
        "Portfolio performance is reported using a flow-adjusted, time-weighted methodology. "
        "External cash flows are treated as occurring at the start of the day, and the history is broken into segments "
        "between flows. Returns for these segments are chained to produce cumulative TWR over each horizon, consistent "
        "with institutional and GIPS-style practice."
    )

    # --- Security & Asset-Class Returns (Modified Dietz) ---
    p = doc.add_paragraph()
    r = p.add_run("Security and asset-class returns (Modified Dietz). ")
    r.bold = True
    p.add_run(
        "For individual securities and asset classes, the report uses a Modified Dietz money-weighted approach. "
        "All ticker-level cash flows (buys and sells) and prices over the horizon are included, and returns are "
        "computed relative to an average invested capital base. Asset-class returns are formed as value-weighted "
        "averages of the underlying security-level returns."
    )

    # --- Profit and Loss (P/L) ---
    p = doc.add_paragraph()
    r = p.add_run("Profit and loss (P/L). ")
    r.bold = True
    p.add_run(
        "Economic P/L is defined consistently throughout the report as: MV_end minus MV_start minus net cash flows "
        "over the period. At the portfolio level, cash flows are external only (contributions and withdrawals). "
        "At the security and asset-class levels, P/L is based on internal trading flows and aligned to the same "
        "time horizons used for the return calculations."
    )

    # --- Benchmarks & Comparisons ---
    p = doc.add_paragraph()
    r = p.add_run("Benchmarks and comparisons. ")
    r.bold = True
    p.add_run(
        "Benchmark figures (such as broad equity or balanced indices) are shown as simple price returns. "
        "They are rebased to the same start dates as the portfolio for each chart or table so that relative "
        "performance reflects differences in allocation and market exposure rather than timing."
    )

    # --- Data Quality & Limitations ---
    p = doc.add_paragraph()
    r = p.add_run("Data quality and limitations. ")
    r.bold = True
    p.add_run(
        "All results depend on the completeness and accuracy of the underlying holdings, transaction, cash flow, "
        "and price data. Minor differences from broker statements may arise from rounding, pricing gaps, or "
        "timing conventions, but these should not be material to the overall conclusions. The methodology is "
        "designed to be transparent, repeatable, and suitable for executive-level review."
    )

    # Final short design principles list to make it "pop"
    doc.add_paragraph("Core design principles:", style="Normal")
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Separate funding decisions (flows) from investment performance.")
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Use true time-weighted returns at the portfolio level.")
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Use money-weighted, cash-flow-aware math for securities and asset classes.")


    # =============================================================
    # SAVE OUTPUT
    # =============================================================

    DOCX_NAME = f"Portfolio_Performance_Report_{timestamp}.docx"
    PDF_NAME  = f"Portfolio_Performance_Report_{timestamp}.pdf"

    doc.save(DOCX_NAME)

    try:
        convert(DOCX_NAME, PDF_NAME)
    except:
        pass

    print("[OK] Report generated:")
    print("   -", DOCX_NAME)
    print("   -", PDF_NAME)


if __name__ == "__main__":
    build_report()
